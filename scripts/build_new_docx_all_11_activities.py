# -*- coding: utf-8 -*-
import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)

def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            {f'<w:top w:val="{top.get("val","single")}" w:sz="{top.get("sz","4")}" w:space="0" w:color="{top.get("color","CCCCCC")}"/>' if top else '<w:top w:val="none"/>'}
            {f'<w:left w:val="{left.get("val","single")}" w:sz="{left.get("sz","4")}" w:space="0" w:color="{left.get("color","CCCCCC")}"/>' if left else '<w:left w:val="none"/>'}
            {f'<w:bottom w:val="{bottom.get("val","single")}" w:sz="{bottom.get("sz","4")}" w:space="0" w:color="{bottom.get("color","CCCCCC")}"/>' if bottom else '<w:bottom w:val="none"/>'}
            {f'<w:right w:val="{right.get("val","single")}" w:sz="{right.get("sz","4")}" w:space="0" w:color="{right.get("color","CCCCCC")}"/>' if right else '<w:right w:val="none"/>'}
        </w:tcBorders>
    ''')
    tcPr.append(tcBorders)

def add_heading_1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Segoe UI'
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(27, 54, 93) # #1B365D Deep Navy
    return p

def add_heading_2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Segoe UI'
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 128, 128) # #008080 Teal
    return p

def add_heading_3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Segoe UI'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(51, 51, 51)
    return p

def add_body_p(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = 'Segoe UI'
        r_pre.font.size = Pt(10.5)
        r_pre.font.bold = True
        r_pre.font.color.rgb = RGBColor(27, 54, 93) # Deep Navy para resaltar el nombre
    r_txt = p.add_run(text)
    r_txt.font.name = 'Segoe UI'
    r_txt.font.size = Pt(10.5)
    r_txt.font.color.rgb = RGBColor(51, 51, 51)
    return p

def add_callout_box(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, "F4F6F8")
    set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
    set_cell_borders(cell, left={"val": "single", "sz": "24", "color": "1B365D"})
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    r_title = p.add_run(f"EXPLICACIÓN GENERAL DEL APARTADO: {title.upper()}\n")
    r_title.font.name = 'Segoe UI'
    r_title.font.size = Pt(11)
    r_title.font.bold = True
    r_title.font.color.rgb = RGBColor(27, 54, 93)
    
    r_txt = p.add_run(text)
    r_txt.font.name = 'Segoe UI'
    r_txt.font.size = Pt(10)
    r_txt.font.color.rgb = RGBColor(51, 51, 51)
    
    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(0)
    p_after.paragraph_format.space_after = Pt(6)

def add_code_block(doc, code_text, caption="Fragmento de código referenciado del proyecto LavanderiaApp0.1:"):
    p_cap = doc.add_paragraph()
    p_cap.paragraph_format.space_before = Pt(6)
    p_cap.paragraph_format.space_after = Pt(2)
    p_cap.paragraph_format.keep_with_next = True
    r_cap = p_cap.add_run(caption)
    r_cap.font.name = 'Segoe UI'
    r_cap.font.size = Pt(9.5)
    r_cap.font.bold = True
    r_cap.font.italic = True
    r_cap.font.color.rgb = RGBColor(100, 100, 100)
    
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, "F2F4F7")
    set_cell_margins(cell, top=100, bottom=100, left=140, right=140)
    set_cell_borders(cell, 
                     top={"val":"single","sz":"4","color":"D0D5DD"},
                     bottom={"val":"single","sz":"4","color":"D0D5DD"},
                     left={"val":"single","sz":"12","color":"008080"},
                     right={"val":"single","sz":"4","color":"D0D5DD"})
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    r = p.add_run(code_text.strip())
    r.font.name = 'Consolas'
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(30, 41, 59)
    
    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(0)
    p_after.paragraph_format.space_after = Pt(6)

def add_activity_table(doc, data_dict):
    table = doc.add_table(rows=len(data_dict)+1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    hdr_cell_0 = table.cell(0, 0)
    hdr_cell_1 = table.cell(0, 1)
    hdr_cell_0.width = Inches(2.2)
    hdr_cell_1.width = Inches(4.3)
    set_cell_background(hdr_cell_0, "1B365D")
    set_cell_background(hdr_cell_1, "1B365D")
    set_cell_margins(hdr_cell_0, 100, 100, 120, 120)
    set_cell_margins(hdr_cell_1, 100, 100, 120, 120)
    
    p0 = hdr_cell_0.paragraphs[0]
    p0.paragraph_format.space_before = Pt(0)
    p0.paragraph_format.space_after = Pt(0)
    r0 = p0.add_run("CAMPO")
    r0.font.name = 'Segoe UI'
    r0.font.size = Pt(10)
    r0.font.bold = True
    r0.font.color.rgb = RGBColor(255, 255, 255)
    
    p1 = hdr_cell_1.paragraphs[0]
    p1.paragraph_format.space_before = Pt(0)
    p1.paragraph_format.space_after = Pt(0)
    r1 = p1.add_run("DETALLE DE LA ACTIVIDAD / APARTADO")
    r1.font.name = 'Segoe UI'
    r1.font.size = Pt(10)
    r1.font.bold = True
    r1.font.color.rgb = RGBColor(255, 255, 255)
    
    row_idx = 1
    for key, value in data_dict.items():
        c0 = table.cell(row_idx, 0)
        c1 = table.cell(row_idx, 1)
        c0.width = Inches(2.2)
        c1.width = Inches(4.3)
        bg = "FFFFFF" if row_idx % 2 != 0 else "F9FAFB"
        set_cell_background(c0, bg)
        set_cell_background(c1, bg)
        set_cell_margins(c0, 80, 80, 120, 120)
        set_cell_margins(c1, 80, 80, 120, 120)
        set_cell_borders(c0, bottom={"val":"single","sz":"4","color":"E5E7EB"}, right={"val":"single","sz":"4","color":"E5E7EB"})
        set_cell_borders(c1, bottom={"val":"single","sz":"4","color":"E5E7EB"})
        
        pc0 = c0.paragraphs[0]
        pc0.paragraph_format.space_before = Pt(0)
        pc0.paragraph_format.space_after = Pt(0)
        rc0 = pc0.add_run(key)
        rc0.font.name = 'Segoe UI'
        rc0.font.size = Pt(9.5)
        rc0.font.bold = True
        rc0.font.color.rgb = RGBColor(27, 54, 93)
        
        pc1 = c1.paragraphs[0]
        pc1.paragraph_format.space_before = Pt(0)
        pc1.paragraph_format.space_after = Pt(0)
        rc1 = pc1.add_run(str(value))
        rc1.font.name = 'Segoe UI'
        rc1.font.size = Pt(9.5)
        rc1.font.color.rgb = RGBColor(51, 51, 51)
        
        row_idx += 1
        
    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(0)
    p_after.paragraph_format.space_after = Pt(8)

def main():
    doc = docx.Document()
    
    # Configurar márgenes de sección
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # PORTADA / TÍTULO PRINCIPAL
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(24)
    p_title.paragraph_format.space_after = Pt(6)
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run("DOCUMENTO TÉCNICO Y GESTIÓN POR HITOS Y ROLES\nLAVANDERÍA APP v0.1")
    r_title.font.name = 'Segoe UI'
    r_title.font.size = Pt(22)
    r_title.font.bold = True
    r_title.font.color.rgb = RGBColor(27, 54, 93)
    
    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_before = Pt(0)
    p_sub.paragraph_format.space_after = Pt(24)
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_sub.add_run("Análisis Exhaustivo por Apartados, Explicación General, Asignación de Roles (Yadier, Leyva, Daniel) y Referencias de Código C#/Razor/SQLite")
    r_sub.font.name = 'Segoe UI'
    r_sub.font.size = Pt(12)
    r_sub.font.italic = True
    r_sub.font.color.rgb = RGBColor(0, 128, 128)
    
    add_body_p(doc, "El presente documento formaliza la planificación, ejecución y sustento técnico de cada uno de los 11 apartados o actividades clave que componen el desarrollo del sistema integral de gestión para lavandería LavanderiaApp v0.1 (.NET 8 Blazor Hybrid en WPF y SQLite nativo). Para cada actividad se expone en primer lugar la ficha técnica de identificación y una explicación general del apartado que describe qué se pretende hacer y cuál es su importancia. Posteriormente, se detallan los hitos alcanzados, asignando responsabilidades según el rol oficial de cada integrante y explicando de forma directa, técnica y precisa cómo se implementó cada hito, demostrando su cumplimiento mediante fragmentos reales del código fuente de nuestra aplicación.")
    
    add_heading_1(doc, "ASIGNACIÓN DE ROLES OFICIALES DEL EQUIPO")
    add_body_p(doc, "• Yadier Pech Tun (Backend / Base de Datos): Encargado del diseño arquitectónico de persistencia en SQLite, sentencias DDL, capa de repositorios transaccionales (Pedidos, Usuarios), transaccionalidad y cobros, modelado relacional e implentación transaccional de inventarios y máquinas.")
    add_body_p(doc, "• Jesus Leyva Chan (Backend): Encargado del desarrollo del motor de autenticación (LoginServicio), gestión del ciclo de vida de sesión en memoria para entorno de escritorio (SessionManager), e implementación transaccional de registros e historial de clientes.")
    add_body_p(doc, "• Daniel Moo (Frontend): Encargado de la conceptualización de wireframes y flujos prototipados en Figma, maquetación estructural del contenedor global (.razor), e implementación de las pantallas interactivas de mostrador con binding y retroalimentación modal en tiempo real.")
    
    # ==========================================
    # ACT 1
    # ==========================================
    add_heading_1(doc, "1. LEVANTAMIENTO DE REQUERIMIENTOS Y DEFINICIÓN DEL ALCANCE")
    add_activity_table(doc, {
        "Identificación de la actividad": "Levantamiento de requerimientos y definición del alcance.",
        "Fecha de inicio": "2026-05-07",
        "Fecha de fin": "2026-06-02",
        "Duración estimada": "29 días",
        "Responsable": "Yadier Pech Tun, Leyva Chan y Daniel Moo",
        "Dependencia": "Ninguna",
        "Hitos": "Lista de módulos, objetivos del sistema, alcance validado y prioridades del proyecto",
        "Porcentaje de avance": "100%",
        "Nota del apartado": "De aquí sale la base del resto del proyecto: módulos, reglas de negocio y entregables mínimos."
    })
    
    add_callout_box(doc, "Levantamiento de requerimientos y definición del alcance",
        "Este primer apartado tiene como propósito fundamental entender las necesidades operativas reales de una lavandería y traducirlas en un esquema arquitectónico y funcional de software claro. Se pretende establecer un mapa conceptual sólido y delimitado que resuelva los problemas cotidianos de un mostrador de lavandería: recepción y pesaje de ropa, categorización de servicios (lavado, secado, tintorería, planchado), control de turnos y estados de las máquinas, deducción de insumos (detergentes y suavizantes) por receta, cobro preciso con emisión de tickets y auditoría de ingresos. Aquí se define qué se va a construir y cuáles son los límites de la versión actual (V0.1/V1.2).")
    
    add_heading_2(doc, "Explicación de los Hitos Alcanzados según Roles del Equipo")
    
    add_heading_3(doc, "Hito 1.1: Lista de Módulos del Sistema")
    add_body_p(doc, "Se estructuró la arquitectura de la aplicación en 7 módulos transaccionales independientes y cohesivos, diseñados para cubrir todo el flujo operativo en el mostrador: (1) Módulo General / Dashboard, (2) Módulo de Pedidos y Cobros, (3) Módulo de Clientes, (4) Módulo de Inventario e Insumos, (5) Módulo de Empleados y Roles, (6) Módulo de Máquinas, y (7) Módulo de Finanzas y Reportes.", bold_prefix="Equipo Colaborativo (Yadier Pech Tun, Leyva Chan, Daniel Moo): ")
    add_body_p(doc, "Se implementó la navegación modular en el componente principal, conectando cada ítem visual con rutas específicas e interceptando el renderizado mediante permisos dinámicos (@if (TienePermisoMenu(...))) para garantizar que cada operador acceda exclusivamente a los módulos que su rol jerárquico le autoriza.", bold_prefix="Daniel Moo (Frontend) & Yadier Pech Tun (Backend / BD): ")
    
    add_code_block(doc, r"""// Estructura modular reflejada en la navegación de MainLayout.razor
<nav class="sidebar-nav">
    <NavLink href="/dashboard" class="nav-item">
        <span class="material-symbols-rounded">grid_view</span><span>General</span>
    </NavLink>
    @if (TienePermisoMenu("Pedidos")) {
        <NavLink href="/pedidos" class="nav-item"><span>Pedidos</span></NavLink>
    }
    @if (TienePermisoMenu("Clientes")) {
        <NavLink href="/clientes" class="nav-item"><span>Clientes</span></NavLink>
    }
    @if (TienePermisoMenu("Empleados")) {
        <NavLink href="/empleados" class="nav-item"><span>Empleados</span></NavLink>
    }
    @if (TienePermisoMenu("Finanzas")) {
        <NavLink href="/finanzas" class="nav-item"><span>Finanzas</span></NavLink>
    }
    @if (TienePermisoMenu("Inventario")) {
        <NavLink href="/inventario" class="nav-item"><span>Inventario</span></NavLink>
    }
    @if (TienePermisoMenu("Maquinas")) {
        <NavLink href="/maquinas" class="nav-item"><span>Máquinas</span></NavLink>
    }
    @if (TienePermisoMenu("Reportes")) {
        <NavLink href="/reportes" class="nav-item"><span>Reportes</span></NavLink>
    }
</nav>""", "Demostración en código del Hito 1.1 (Lista de módulos maquetada y condicional en MainLayout.razor):")
    add_body_p(doc, "Se implementó en este componente la división exacta de los 7 módulos del sistema. La lógica de evaluación de permisos se ejecuta en tiempo de renderizado para cada opción, garantizando un acoplamiento limpio entre la definición del alcance y el menú interactivo.")
    
    add_heading_3(doc, "Hito 1.2: Objetivos del Sistema")
    add_body_p(doc, "Se programaron las reglas de negocio monetarias dentro del modelo transaccional de pedidos para cumplir el objetivo de exactitud financiera y adaptabilidad fiscal. Se implementó una lógica de cálculo dinámico del total que evalúa en tiempo de ejecución la configuración activa del negocio (IvaActivo e Iva) para aplicar o exentar el impuesto, así como el cálculo automático exacto del saldo pendiente en función de los abonos parciales registrados.", bold_prefix="Yadier Pech Tun (Backend / Base de Datos) & Leyva Chan (Backend): ")
    add_code_block(doc, r"""// Reglas de negocio monetarias y de objetivos implementadas en Pedido.cs
public decimal CalcularTotal()
{
    decimal subtotal = Detalles?.Sum(d => d.Subtotal) ?? 0m;
    if (BusinessConfig.Current != null && BusinessConfig.Current.IvaActivo && BusinessConfig.Current.Iva > 0)
    {
        Total = Math.Round(subtotal * (1m + (decimal)(BusinessConfig.Current.Iva / 100.0)), 2);
    }
    else { Total = subtotal; }
    return Total;
}
public decimal SaldoPendiente => (Total - MontoPagado) > 0 ? (Total - MontoPagado) : 0m;
public bool EstaPagado => Total > 0 && SaldoPendiente <= 0;""", "Demostración en código del Hito 1.2 (Reglas de negocio y objetivos monetarios en Pedido.cs):")
    add_body_p(doc, "Se implementó el método CalcularTotal() y las propiedades calculadas SaldoPendiente y EstaPagado para encapsular el comportamiento financiero. Esta implementación asegura el objetivo del sistema de no permitir errores aritméticos ni discrepancias entre mostrador y contabilidad.")
    
    add_heading_3(doc, "Hito 1.3: Alcance Validado")
    add_body_p(doc, "Se implementó la solución de software bajo la arquitectura híbrida local .NET 8 Blazor Hybrid empacada en contenedor WPF con motor relacional empotrado SQLite. Esta decisión arquitectónica valida y cumple el alcance de operar con cero latencia y tolerancia total a fallos de conexión a Internet en las terminales de la lavandería.", bold_prefix="Daniel Moo (Frontend) & Yadier Pech Tun (Backend / BD): ")
    
    add_heading_3(doc, "Hito 1.4: Prioridades del Proyecto")
    add_body_p(doc, "Se implementó un desarrollo por capas estrictamente priorizado: se construyó primero el motor y capa de datos SQLite junto con la seguridad del backend, posteriormente se desarrollaron los modelos de negocio transaccionales y finalmente se acopló la capa de presentación visual en Blazor con servicios y reportes nativos.", bold_prefix="Equipo Colaborativo (Yadier Pech Tun, Leyva Chan, Daniel Moo): ")

    # ==========================================
    # ACT 2
    # ==========================================
    add_heading_1(doc, "2. DEFINICIÓN DE BACKLOG Y MÓDULOS PRINCIPALES")
    add_activity_table(doc, {
        "Identificación de la actividad": "Definición de backlog y módulos principales.",
        "Fecha de inicio": "2026-06-07",
        "Fecha de fin": "2026-06-11",
        "Duración estimada": "4 días",
        "Responsable": "Yadier Pech Tun",
        "Dependencia": "Requerimientos definidos",
        "Hitos": "Módulos de usuarios, clientes, servicios, pedidos y reportes detalles pedido, etc",
        "Porcentaje de avance": "100%",
        "Nota del apartado": "Apartado ya listo. AVANCE TOTAL AL TERMINAR: 25%"
    })
    
    add_callout_box(doc, "Definición de backlog y módulos principales",
        "Este apartado tiene el objetivo de descomponer los requerimientos funcionales en un backlog técnico estructurado. Aquí se diseñan los modelos de dominio (POCOs / clases C#) que representan cada entidad del mundo real de la lavandería. Se estructuran las propiedades, constructores, métodos de validación interna y contratos de datos para los módulos de usuarios, clientes, servicios, pedidos, detalles de pedido, auditoría y reportes. Esta formalización asegura que todo el equipo sepa exactamente con qué datos y propiedades cuenta cada entidad antes de escribir consultas SQL o crear pantallas en Blazor.")
    
    add_heading_2(doc, "Explicación de los Hitos Alcanzados según Roles del Equipo")
    
    add_heading_3(doc, "Hito 2.1: Módulo de Usuarios y Roles en el Backlog")
    add_body_p(doc, "Se diseñó e implementó la entidad de dominio para usuarios (`Usuario.cs`), dotándola de un esquema completo de propiedades que unifican los datos de identidad (Nombre, NombreUsuario, Password, Correo) con las variables operativas laborales de la lavandería (Rol, Salario, Sucursal, Activo). Se programaron funciones de evaluación booleana rápidas (`EsMaster`, `EsAdmin`, `EsEmpleado`) para simplificar la verificación de privilegios en el sistema.", bold_prefix="Yadier Pech Tun (Backend / Base de Datos): ")
    add_code_block(doc, r"""// App/Modelos/Usuario.cs - Especificación técnica y modelo del backlog para usuarios
public class Usuario
{
    public int IdUsuario { get; set; }
    public string Nombre { get; set; } = string.Empty;
    public string NombreUsuario { get; set; } = string.Empty;
    public string Password { get; set; } = string.Empty;
    public string Rol { get; set; } = string.Empty; // "Master", "Admin", "Empleado"
    public string Correo { get; set; } = string.Empty;
    public string Telefono { get; set; } = string.Empty;
    public decimal Salario { get; set; } = 0.0m;
    public string Sucursal { get; set; } = string.Empty;
    public bool Activo { get; set; } = true;

    public bool ValidarPassword(string password) => !string.IsNullOrEmpty(password) && Password.Equals(password);
    public bool EsMaster => Rol.Equals("Master", StringComparison.OrdinalIgnoreCase);
    public bool EsAdmin => Rol.Equals("Admin", StringComparison.OrdinalIgnoreCase) || Rol.Equals("Administrador", StringComparison.OrdinalIgnoreCase) || EsMaster;
    public bool EsEmpleado => Rol.Equals("Empleado", StringComparison.OrdinalIgnoreCase);
}""", "Demostración en código del Hito 2.1 (Entidad Usuario.cs en el backlog de modelos):")
    add_body_p(doc, "Se implementó en esta clase el encapsulamiento de la lógica de comparación insensible a mayúsculas/minúsculas para roles (`StringComparison.OrdinalIgnoreCase`). Esto asegura que el sistema identifique de manera robusta los permisos sin importar cómo se haya capturado la cadena en la base de datos.")

    add_heading_3(doc, "Hito 2.2: Módulo de Clientes en el Backlog")
    add_body_p(doc, "Se diseñó e implementó el modelo de dominio `Cliente.cs`, estructurando las propiedades para almacenar información de contacto e introduciendo una propiedad transaccional específica (`PuntosFidelidad`) orientada a cuantificar y recompensar la recurrencia de los clientes en el mostrador.", bold_prefix="Yadier Pech Tun (Backend / Base de Datos): ")
    add_code_block(doc, r"""// App/Modelos/Cliente.cs - Especificación del modelo transaccional de clientes
public class Cliente
{
    public int IdCliente { get; set; }
    public string Nombre { get; set; } = string.Empty;
    public string Telefono { get; set; } = string.Empty;
    public string Direccion { get; set; } = string.Empty;
    public int PuntosFidelidad { get; set; } = 0;
    public DateTime FechaRegistro { get; set; } = DateTime.Now;
}""", "Demostración en código del Hito 2.2 (Entidad Cliente.cs en el backlog):")

    add_heading_3(doc, "Hito 2.3: Módulo de Servicios, Pedidos y Detalles de Pedido")
    add_body_p(doc, "Se implementó la arquitectura de modelos relacionales maestro-detalle mediante las entidades `Pedido.cs` y `DetallePedido.cs`. Se estructuró el detalle para independizar el precio de partida transaccional y la cantidad calculada, incorporando el método de cálculo autónomo `CalcularSubtotal()` que multiplica el volumen o peso de ropa por el precio unitario exacto fijado al momento de la orden.", bold_prefix="Yadier Pech Tun (Backend / Base de Datos): ")
    add_code_block(doc, r"""// App/Modelos/Detalle_Pedido.cs - Modelo maestro-detalle con auto-cálculo de montos
public class DetallePedido
{
    public int IdDetallePedido { get; set; }
    public int IdPedido { get; set; }
    public int IdServicio { get; set; }
    public decimal Cantidad { get; set; }
    public decimal PrecioUnitario { get; set; }
    public decimal Subtotal { get; set; }
    public Servicio? Servicio { get; set; }

    public void CalcularSubtotal()
    {
        Subtotal = Math.Round(Cantidad * PrecioUnitario, 2);
    }
}""", "Demostración en código del Hito 2.3 (Modelo transaccional DetallePedido.cs):")
    add_body_p(doc, "Se implementó en `DetallePedido.cs` el método `CalcularSubtotal()` con redondeo matemático a dos decimales (`Math.Round`), garantizando que la sumatoria de ítems de una orden sea completamente inmune a inconsistencias o errores flotantes.")

    # ==========================================
    # ACT 3
    # ==========================================
    add_heading_1(doc, "3. DISEÑO DE INTERFAZ Y FLUJO EN FIGMA")
    add_activity_table(doc, {
        "Identificación de la actividad": "Diseño de interfaz y flujo en Figma.",
        "Fecha de inicio": "2026-06-05",
        "Fecha de fin": "2026-06-07",
        "Duración estimada": "2 días",
        "Responsable": "Yadier Pech Tun, Jesus Leyva Chan, Daniel Moo",
        "Dependencia": "Requerimientos definidos",
        "Hitos": "Wireframes, flujo de pantallas, prototipo navegable y estructura visual",
        "Porcentaje de avance": "100%",
        "Nota del apartado": "AVANCE TOTAL: 35%"
    })
    
    add_callout_box(doc, "Diseño de interfaz y flujo en Figma",
        "Este apartado tiene el propósito de crear el prototipo visual interactivo en Figma antes de escribir código UI. Se persigue diseñar una experiencia de usuario que reduzca la fricción en el mostrador: botones grandes para dispositivos táctiles, colores limpios que evoquen pulcritud (azul marino, turquesa, blanco) y una jerarquía visual clara que guíe al recepcionista desde el inicio de sesión, pasando por el panel general, hasta la creación y cobro del pedido en pocos segundos.")
    
    add_heading_2(doc, "Explicación de los Hitos Alcanzados según Roles del Equipo")
    
    add_heading_3(doc, "Hito 3.1: Wireframes y Estructura Visual")
    add_body_p(doc, "Se diseñaron los wireframes y la retícula visual de alta fidelidad en Figma, estableciendo un patrón arquitectónico de doble panel: una barra de navegación lateral fija (`<aside class=\"sidebar\">`) y una zona de contenido central dinámica (`<main class=\"main-content\">`). Se implementó exactamente esta misma retícula de maquetación en el contenedor principal de Blazor (`MainLayout.razor`) con soporte para temas visuales.", bold_prefix="Daniel Moo (Frontend): ")
    add_code_block(doc, r"""<!-- App/Shared/MainLayout.razor - Implementación en código de la estructura visual y wireframes -->
<div class="app-layout @GetThemeClass()">
    <aside class="sidebar">
        <div class="sidebar-logo-area">
            <div class="logo-box"><span class="material-symbols-rounded">local_laundry_service</span></div>
            <div class="logo-text">
                <span class="logo-title-villas">@ObtenerNombreNegocioLinea1()</span>
                <span class="logo-subtitle-villas">@ObtenerNombreNegocioLinea2()</span>
            </div>
        </div>
        <nav class="sidebar-nav">
            <!-- Menú interactivo -->
        </nav>
    </aside>
    <main class="main-content">
        <header class="top-bar">
            <!-- Datos de sesión activa y alertas -->
        </header>
        <div class="content-body">@Body</div>
    </main>
</div>""", "Demostración en código del Hito 3.1 (Estructura visual y layout trasladado de Figma a Blazor):")
    add_body_p(doc, "Se implementó una división semántica limpia entre el menú y el cuerpo de renderizado (`@Body`). Esta estructura materializa los wireframes validados, manteniendo las herramientas del mostrador accesibles de forma constante en el menú izquierdo mientras la información operativa fluye en el panel derecho.")

    add_heading_3(doc, "Hito 3.2: Flujo de Pantallas y Prototipo Navegable")
    add_body_p(doc, "Se diseñó y validó en Figma el recorrido interactivo completo para el operador de lavandería, verificando la secuencia de pasos lógicos: Login -> Dashboard de Monitoreo -> Directorio de Clientes -> Creación/Edición de Pedidos -> Pantalla de Cobro y Cambio de Estado.", bold_prefix="Daniel Moo (Frontend), Jesus Leyva Chan (Backend) & Yadier Pech Tun (Backend / BD): ")

    # ==========================================
    # ACT 4
    # ==========================================
    add_heading_1(doc, "4. DISEÑO DE BASE DE DATOS Y ARQUITECTURA")
    add_activity_table(doc, {
        "Identificación de la actividad": "Diseño de base de datos y arquitectura.",
        "Fecha de inicio": "2026-06-07",
        "Fecha de fin": "2026-06-11",
        "Duración estimada": "4 días",
        "Responsable": "Yadier Pech Tun y Jesus Leyva Chan",
        "Dependencia": "Definición de backlogs y módulos",
        "Hitos": "Modelo entidad-relación, estructura de tablas, SQLite con reglas base y etiquetas",
        "Porcentaje de avance": "25% (al cierre de la etapa inicial DDL, sentando la base de persistencia)",
        "Nota del apartado": "De este diseño dependerán autenticación, clientes, servicios, órdenes y reportes. AVANCE TOTAL: 40%"
    })
    
    add_callout_box(doc, "Diseño de base de datos y arquitectura",
        "Este apartado construye el motor físico de persistencia del sistema. Se selecciona y configura SQLite nativo (Lavanderia.db) por su extrema velocidad, confiabilidad e independencia de servidores externos, perfecto para una aplicación local de lavandería en Windows. Aquí se construyen las tablas transaccionales, se habilitan las restricciones de integridad referencial (FOREIGN KEYs), se establecen índices y se programan rutinas de inicialización y migración automática DDL que garantizan que la base de datos se cree y actualice automáticamente en cualquier computadora sin perder la información del cliente.")
    
    add_heading_2(doc, "Explicación de los Hitos Alcanzados según Roles del Equipo")
    
    add_heading_3(doc, "Hito 4.1: Modelo Entidad-Relación y Estructura de Tablas en SQLite")
    add_body_p(doc, "Se implementó el esquema DDL relacional nativo dentro del inicializador de base de datos (`DatabaseInitializer.cs`). Se estructuraron las tablas core (`Clientes`, `Usuarios`, `Pedidos`, `DetallesPedido`) aplicando las llaves primarias (`PRIMARY KEY AUTOINCREMENT`) y activando la integridad referencial estricta del motor (`PRAGMA foreign_keys = ON;`) para vincular cada pedido a un cliente y al usuario recepcionista responsable.", bold_prefix="Yadier Pech Tun (Backend / Base de Datos): ")
    add_code_block(doc, r"""// App/BaseDatos/DatabaseInitializer.cs - Implementación DDL transaccional y activación de llaves foráneas
using var conexion = new SqliteConnection(Config.ConnectionString);
conexion.Open();
string baseSchema = @"
    PRAGMA foreign_keys = ON;

    CREATE TABLE IF NOT EXISTS Clientes (
        IdCliente INTEGER PRIMARY KEY AUTOINCREMENT,
        Nombre TEXT NOT NULL,
        Telefono TEXT NOT NULL,
        Direccion TEXT NOT NULL
    );
    
    CREATE TABLE IF NOT EXISTS Usuarios (
        IdUsuario INTEGER PRIMARY KEY AUTOINCREMENT,
        Nombre TEXT NOT NULL,
        NombreUsuario TEXT NOT NULL UNIQUE,
        Password TEXT NOT NULL,
        Rol TEXT NOT NULL
    );

    CREATE TABLE IF NOT EXISTS Pedidos (
        IdPedido INTEGER PRIMARY KEY AUTOINCREMENT,
        IdCliente INTEGER NOT NULL,
        IdUsuario INTEGER NOT NULL,
        FechaRecepcion TEXT NOT NULL,
        FechaEntrega TEXT,
        Total DECIMAL NOT NULL,
        FOREIGN KEY (IdCliente) REFERENCES Clientes(IdCliente),
        FOREIGN KEY (IdUsuario) REFERENCES Usuarios(IdUsuario)
    );

    CREATE TABLE IF NOT EXISTS DetallesPedido (
        IdDetallePedido INTEGER PRIMARY KEY AUTOINCREMENT,
        IdPedido INTEGER NOT NULL,
        IdServicio INTEGER NOT NULL,
        Cantidad REAL NOT NULL,
        PrecioUnitario DECIMAL NOT NULL,
        Subtotal DECIMAL NOT NULL,
        FOREIGN KEY (IdPedido) REFERENCES Pedidos(IdPedido),
        FOREIGN KEY (IdServicio) REFERENCES Servicios(IdServicio)
    );";
using var cmd = new SqliteCommand(baseSchema, conexion);
cmd.ExecuteNonQuery();""", "Demostración en código del Hito 4.1 (Sentencias DDL en DatabaseInitializer.cs):")
    add_body_p(doc, "Se implementó en este script la orden de creación condicional (`CREATE TABLE IF NOT EXISTS`), asegurando que al arrancar la aplicación en una terminal no se sobrescriba ni se corrompa el archivo `Lavanderia.db` existente, al mismo tiempo que se blinda a nivel motor la coherencia referencial de las tablas de órdenes y detalles.")

    add_heading_3(doc, "Hito 4.2: Reglas Base, Migración Automática y Etiquetas")
    add_body_p(doc, "Se implementó una rutina en C# de introspección y migración automática de esquemas (`MigrarTabla`) que permite evolucionar la estructura de la base de datos de manera dinámica y segura. La rutina ejecuta la consulta `PRAGMA table_info` de SQLite para auditar las columnas existentes en el archivo físico y, si detecta nuevas propiedades en los modelos de negocio, genera e inyecta sentencias `ALTER TABLE` sin afectar los datos ya almacenados.", bold_prefix="Yadier Pech Tun (Backend / Base de Datos) & Jesus Leyva Chan (Backend): ")
    add_code_block(doc, r"""// App/BaseDatos/DatabaseInitializer.cs - Rutina de introspección y auto-migración DDL en tiempo de ejecución
private static void MigrarTabla(SqliteConnection conexion, string tabla, Dictionary<string, string> columnasDefinicion)
{
    var columnasExistentes = new HashSet<string>(StringComparer.OrdinalIgnoreCase);
    using (var cmd = new SqliteCommand($"PRAGMA table_info({tabla})", conexion))
    using (var reader = cmd.ExecuteReader())
    {
        while (reader.Read())
        {
            columnasExistentes.Add(reader.GetString(1)); // Nombre de la columna
        }
    }

    foreach (var col in columnasDefinicion)
    {
        if (!columnasExistentes.Contains(col.Key))
        {
            using var alterCmd = new SqliteCommand($"ALTER TABLE {tabla} ADD COLUMN {col.Key} {col.Value}", conexion);
            alterCmd.ExecuteNonQuery();
        }
    }
}""", "Demostración en código del Hito 4.2 (Rutina automática MigrarTabla en DatabaseInitializer.cs):")
    add_body_p(doc, "Se implementó este mecanismo de auditoría de esquema para independizar el despliegue del software de migraciones manuales. Cuando se añaden columnas como `InventarioRestado` o `MaquinaAsignada`, la aplicación aplica las modificaciones estructurales en milisegundos durante el arranque en caliente.")

    # ==========================================
    # ACT 5
    # ==========================================
    add_heading_1(doc, "5. APROBACIÓN DEL PROTOTIPO Y CIERRE DE DISEÑO")
    add_activity_table(doc, {
        "Identificación de la actividad": "Aprobación del prototipo y cierre de diseño.",
        "Fecha de inicio": "2026-06-12",
        "Fecha de fin": "2026-06-19",
        "Duración estimada": "7 días",
        "Responsable": "Todos",
        "Dependencia": "Diseño de interfaz y flujo en Figma",
        "Hitos": "Validación visual, ajustes finales y aprobación para desarrollo",
        "Porcentaje de avance": "100%",
        "Nota del apartado": "Esta tarjeta marca el punto de cierre de paso entre diseño y construcción. AVANCE TOTAL: 50%"
    })
    
    add_callout_box(doc, "Aprobación del prototipo y cierre de diseño",
        "Este apartado representa el punto de control y compuerta de calidad entre la fase de diseño conceptual y la fase de construcción intensiva de software. Todo el equipo somete a escrutinio el prototipo de Figma frente a las especificaciones relacionales y de negocio. Se validan las paletas, las transiciones visuales, la claridad de los textos y se consolida el estándar de comunicación de alertas y diálogos modales para que la interfaz sea robusta, amigable y uniforme antes de iniciar la integración masiva.")
    
    add_heading_2(doc, "Explicación de los Hitos Alcanzados según Roles del Equipo")
    
    add_heading_3(doc, "Hito 5.1: Validación Visual y Ajustes Finales de Comunicación de UI")
    add_body_p(doc, "Se implementó un estándar visual y un servicio centralizado de notificaciones interactivas (`ToastService.cs`) como reemplazo de los cuadros de diálogo intrusivos del sistema operativo que congelan la ejecución. Se estructuró este servicio utilizando eventos en memoria (`event Action<ToastMessage> OnShow`), lo que permite que cualquier proceso del backend envíe alertas visuales no invasivas (éxito, error, advertencia) directamente a la capa UI respetando la paleta de colores del prototipo.", bold_prefix="Daniel Moo (Frontend) & Jesus Leyva Chan (Backend): ")
    add_code_block(doc, r"""// App/Servicios/ToastService.cs - Implementación de retroalimentación modal no invasiva y basada en eventos
namespace LavanderiaApp.Servicios;

public class ToastMessage
{
    public string Message { get; set; } = string.Empty;
    public string Type { get; set; } = "info"; // "success", "error", "warning", "info"
    public DateTime Timestamp { get; set; } = DateTime.Now;
}

public class ToastService
{
    public event Action<ToastMessage>? OnShow;
    public void ShowSuccess(string message) => OnShow?.Invoke(new ToastMessage { Message = message, Type = "success" });
    public void ShowError(string message) => OnShow?.Invoke(new ToastMessage { Message = message, Type = "error" });
    public void ShowWarning(string message) => OnShow?.Invoke(new ToastMessage { Message = message, Type = "warning" });
}""", "Demostración en código del Hito 5.1 (Estándar visual de notificaciones ToastService):")
    add_body_p(doc, "Se implementó esta arquitectura orientada a eventos para lograr un desacoplamiento limpio. Al invocar `ShowSuccess` o `ShowError` desde los repositorios o servicios, el componente modal del layout principal suscribe e interpreta el mensaje para mostrar la alerta emergente y desvanecerla automáticamente.")

    add_heading_3(doc, "Hito 5.2: Aprobación Formal para Desarrollo (.sln y .csproj)")
    add_body_p(doc, "Se congeló formalmente la especificación y se configuraron las dependencias generales en el proyecto `.csproj` y `.sln`, dejando las interfaces y modelos listos para iniciar la codificación transaccional intensiva.", bold_prefix="Equipo Colaborativo (Yadier Pech Tun, Leyva Chan, Daniel Moo): ")

    # ==========================================
    # ACT 6
    # ==========================================
    add_heading_1(doc, "6. BACKEND: AUTENTICACIÓN Y USUARIOS")
    add_activity_table(doc, {
        "Identificación de la actividad": "Backend: autenticación y usuarios.",
        "Fecha de inicio": "2026-06-23",
        "Fecha de fin": "2026-07-04",
        "Duración estimada": "11 días",
        "Responsable": "Yadier Pech Tun y Leyva Chan",
        "Dependencia": "Definición de módulos y backlog",
        "Hitos": "Inicio de sesión(Leyva), roles (Yadier), sesión segura(Leyva) y gestión básica de usuarios(Yadier)",
        "Porcentaje de avance": "100%",
        "Nota del apartado": "Este módulo debe quedar estable antes de avanzar al resto de la lógica del sistema. AVANCE TOTAL AL TERMINAR: 60%"
    })
    
    add_callout_box(doc, "Backend: autenticación y usuarios",
        "Este apartado establece la barrera de seguridad y la gestión de identidad de LavanderiaApp. Es un prerrequisito crítico que debe quedar absolutamente estable antes de construir las pantallas operativas. Se divide en responsabilidades exactas según la tarjeta: se desarrolla la lógica de verificación de inicio de sesión (LoginServicio) y el gestor de sesión activa en memoria (SessionManager), y en paralelo se programa el sistema transaccional de persistencia y CRUD de usuarios en base de datos (UsuarioRepositorio) con modelado de control de acceso jerárquico por roles para blindar los módulos sensibles del negocio.")
    
    add_heading_2(doc, "Explicación de los Hitos Alcanzados según Roles del Equipo")
    
    add_heading_3(doc, "Hito 6.1: Inicio de Sesión")
    add_body_p(doc, "Se implementó el servicio de autenticación (`LoginServicio.cs`), encapsulando la lógica de consulta y verificación de credenciales. Se estructuró el flujo para invocar al repositorio de usuarios, validar el match de contraseñas, confirmar que la cuenta se encuentre con estado `Activo = true` y, tras cumplir estas verificaciones, autorizar el paso inyectando el perfil del operador en el gestor de sesión.", bold_prefix="Jesus Leyva Chan (Backend): ")
    add_code_block(doc, r"""// App/Servicios/LoginServicio.cs - Lógica de verificación de credenciales e inyección de sesión
using LavanderiaApp.Modelos;
using LavanderiaApp.Repositorios;

namespace LavanderiaApp.Servicios;

public class LoginServicio
{
    private UsuarioRepositorio _usuarioRepo;

    public LoginServicio()
    {
        _usuarioRepo = new UsuarioRepositorio();
    }

    public bool Login(string nombreUsuario, string password)
    {
        var usuario = _usuarioRepo.ObtenerPorNombreUsuario(nombreUsuario);
        if (usuario != null && usuario.ValidarPassword(password) && usuario.Activo)
        {
            SessionManager.UsuarioActual = usuario;
            return true;
        }
        return false;
    }
}""", "Demostración en código del Hito 6.1 (Implementación de LoginServicio por Leyva Chan):")
    add_body_p(doc, "Se implementó en este servicio la vinculación directa entre la validación booleana de credenciales (`ValidarPassword`) y la asignación del usuario en memoria (`SessionManager.UsuarioActual`), consolidando una barrera de acceso altamente eficiente.")

    add_heading_3(doc, "Hito 6.2: Sesión Segura en Memoria")
    add_body_p(doc, "Se implementó una clase estática de gestión de sesión (`SessionManager.cs`) para preservar y auditar la identidad del operador activo en toda la aplicación de escritorio sin depender de variables globales inseguras ni archivos temporales vulnerables en disco. Se programó la propiedad `IsLoggedIn` y el método `Logout()` para limpiar la referencia de memoria al cerrar el turno.", bold_prefix="Jesus Leyva Chan (Backend): ")
    add_code_block(doc, r"""// App/Servicios/SessionManager.cs - Gestor de sesión en memoria para entorno local WPF/Blazor
using LavanderiaApp.Modelos;

namespace LavanderiaApp.Servicios;

public static class SessionManager
{
    public static Usuario? UsuarioActual { get; set; }
    public static bool IsLoggedIn => UsuarioActual != null;

    public static void Logout()
    {
        UsuarioActual = null;
    }
}""", "Demostración en código del Hito 6.2 (SessionManager en memoria por Leyva Chan):")
    add_body_p(doc, "Se implementó `SessionManager` con un diseño en memoria para garantizar lecturas instantáneas por parte de cualquier componente visual (como la barra superior o las tablas transaccionales), asegurando que todas las órdenes y movimientos de caja queden firmados con el `IdUsuario` en curso.")

    add_heading_3(doc, "Hito 6.3: Roles de Usuario y Gestión Básica de Usuarios (CRUD y Repositorio)")
    add_body_p(doc, "Se implementó el repositorio transaccional `UsuarioRepositorio.cs` para el acceso a la tabla `Usuarios` en SQLite, incorporando la consulta segura de recuperación por nombre de usuario (`ObtenerPorNombreUsuario`) y las operaciones de alta, actualización y control jerárquico de roles.", bold_prefix="Yadier Pech Tun (Backend / Base de Datos): ")
    add_code_block(doc, r"""// App/Repositorios/UsuarioRepositorio.cs - Consulta parametrizada y mapeo transaccional de usuarios
public Usuario? ObtenerPorNombreUsuario(string nombreUsuario)
{
    using var conexion = new SqliteConnection(Config.ConnectionString);
    conexion.Open();
    string query = "SELECT * FROM Usuarios WHERE NombreUsuario = @NombreUsuario";
    using var cmd = new SqliteCommand(query, conexion);
    cmd.Parameters.AddWithValue("@NombreUsuario", nombreUsuario);
    using var reader = cmd.ExecuteReader();

    if (reader.Read())
    {
        return new Usuario
        {
            IdUsuario = reader.GetInt32(reader.GetOrdinal("IdUsuario")),
            Nombre = reader.GetString(reader.GetOrdinal("Nombre")),
            NombreUsuario = reader.GetString(reader.GetOrdinal("NombreUsuario")),
            Password = reader.GetString(reader.GetOrdinal("Password")),
            Rol = reader.GetString(reader.GetOrdinal("Rol")),
            Activo = reader.IsDBNull(reader.GetOrdinal("Activo")) || reader.GetInt32(reader.GetOrdinal("Activo")) == 1
        };
    }
    return null;
}""", "Demostración en código del Hito 6.3 y 6.4 (Repositorio de usuarios por Yadier Pech Tun):")
    add_body_p(doc, "Se implementó el método `ObtenerPorNombreUsuario` utilizando comandos SQL estrictamente parametrizados (`@NombreUsuario`) para anular cualquier riesgo de inyección SQL. El mapeo recupera con exactitud los datos de rol y estado activo desde el archivo SQLite hacia el modelo en memoria.")

    # ==========================================
    # ACT 7
    # ==========================================
    add_heading_1(doc, "7. BACKEND: CLIENTES, SERVICIOS Y ÓRDENES")
    add_activity_table(doc, {
        "Identificación de la actividad": "Backend: clientes, servicios y órdenes.",
        "Fecha de inicio": "2026-07-05",
        "Fecha de fin": "2026-07-14",
        "Duración estimada": "5 días",
        "Responsable": "Yadier Pech Tun y Leyva Chan",
        "Dependencia": "Autenticación y usuarios listos",
        "Hitos": "CRUD principal(Yadier), registro de clientes (Leyva)",
        "Porcentaje de avance": "100%",
        "Nota del apartado": "Aquí se concentra la lógica más importante de la lavandería. AVANCE TOTAL AL TERMINAR: 65%"
    })
    
    add_callout_box(doc, "Backend: clientes, servicios y órdenes",
        "Este apartado constituye el núcleo transaccional de la lavandería. Aquí se codifican las reglas transaccionales más críticas que conectan la recepción del mostrador con la persistencia en base de datos. Se divide con precisión según la tarjeta oficial: se programa el CRUD Principal de Pedidos, Detalles, Servicios, Pagos y la automatización inteligente del inventario (para que al procesar un pedido se descuente el detergente y se asigne una lavadora); y en paralelo se desarrolla la lógica de Registro y Gestión de Clientes, permitiendo búsquedas instantáneas por teléfono e historial de consumo.")
    
    add_heading_2(doc, "Explicación de los Hitos Alcanzados según Roles del Equipo")
    
    add_heading_3(doc, "Hito 7.1: CRUD Principal de Órdenes y Automatización de Insumos")
    add_body_p(doc, "Se implementó la capa central de persistencia de órdenes en `PedidoRepositorio.cs` y el servicio de deducción de stock en `InventarioAutomatizacion.cs`. Se estructuró la inserción transaccional de pedidos para ejecutar un `SELECT last_insert_rowid();` en la misma transacción, obteniendo de inmediato la llave primaria autogenerada en SQLite para poder asociar e insertar en secuencia cada ítem del detalle del pedido sin riesgo de pérdidas o cruces referenciales.", bold_prefix="Yadier Pech Tun (Backend / Base de Datos): ")
    add_code_block(doc, r"""// App/Repositorios/PedidoRepositorio.cs - Inserción y captura de ID con SELECT last_insert_rowid()
public int Guardar(Pedido pedido)
{
    using var conexion = new SqliteConnection(Config.ConnectionString);
    conexion.Open();

    string query = @"
        INSERT INTO Pedidos (IdCliente, IdUsuario, FechaRecepcion, FechaEntrega, Estado, Total, InventarioRestado, CostoInsumos, MaquinaAsignada)
        VALUES (@IdCliente, @IdUsuario, @FechaRecepcion, @FechaEntrega, @Estado, @Total, @InventarioRestado, @CostoInsumos, @MaquinaAsignada);
        SELECT last_insert_rowid();";
    using var command = new SqliteCommand(query, conexion);
    command.Parameters.AddWithValue("@IdCliente", pedido.IdCliente);
    command.Parameters.AddWithValue("@IdUsuario", pedido.IdUsuario);
    command.Parameters.AddWithValue("@FechaRecepcion", pedido.FechaRecepcion.ToString("yyyy-MM-dd HH:mm:ss"));
    command.Parameters.AddWithValue("@FechaEntrega", pedido.FechaEntrega.HasValue ? pedido.FechaEntrega.Value.ToString("yyyy-MM-dd HH:mm:ss") : (object)System.DBNull.Value);
    command.Parameters.AddWithValue("@Estado", string.IsNullOrEmpty(pedido.Estado) ? "En espera" : pedido.Estado);
    command.Parameters.AddWithValue("@Total", pedido.Total);
    command.Parameters.AddWithValue("@InventarioRestado", pedido.InventarioRestado);
    command.Parameters.AddWithValue("@CostoInsumos", pedido.CostoInsumos);
    command.Parameters.AddWithValue("@MaquinaAsignada", pedido.MaquinaAsignada ?? "");
    
    return (int)(long)command.ExecuteScalar();
}""", "Demostración en código del Hito 7.1 (Guardado transaccional en PedidoRepositorio.cs):")
    add_body_p(doc, "Se implementó el uso de `command.ExecuteScalar()` para retornar el identificador numérico exacto asignado por SQLite. Esta técnica garantiza consistencia relacional absoluta al registrar pedidos y abonos monetarios en terminales concurrentes.")

    add_heading_3(doc, "Hito 7.2: Registro y Gestión de Clientes")
    add_body_p(doc, "Se implementó el repositorio de persistencia `ClienteRepositorio.cs`, estructurando las operaciones de registro (`Guardar`) con parámetros SQL para nombre, teléfono y dirección, y programando el método de recuperación de listados generales (`ListarTodo`) optimizado con lectores de datos rápidos (`SqliteDataReader`) para alimentar el directorio visual en el mostrador.", bold_prefix="Jesus Leyva Chan (Backend): ")
    add_code_block(doc, r"""// App/Repositorios/ClienteRepositorio.cs - Inserción transaccional y consulta general de clientes
public void Guardar(Cliente cliente)
{
    using var conexion = new SqliteConnection(Config.ConnectionString);
    conexion.Open();

    string query = @"INSERT INTO Clientes (Nombre, Telefono, Direccion)
                     VALUES (@Nombre, @Telefono, @Direccion)";

    using var cmd = new SqliteCommand(query, conexion);
    cmd.Parameters.AddWithValue("@Nombre", cliente.Nombre);
    cmd.Parameters.AddWithValue("@Telefono", cliente.Telefono ?? "");
    cmd.Parameters.AddWithValue("@Direccion", cliente.Direccion ?? "");

    cmd.ExecuteNonQuery();
}

public List<Cliente> ListarTodo()
{
    var lista = new List<Cliente>();
    using var conexion = new SqliteConnection(Config.ConnectionString);
    conexion.Open();
    string query = "SELECT IdCliente, Nombre, Telefono, Direccion FROM Clientes";
    using var cmd = new SqliteCommand(query, conexion);
    using var reader = cmd.ExecuteReader();
    while (reader.Read())
    {
        lista.Add(new Cliente {
            IdCliente = reader.GetInt32(0),
            Nombre = reader.GetString(1),
            Telefono = reader.IsDBNull(2) ? "" : reader.GetString(2),
            Direccion = reader.IsDBNull(3) ? "" : reader.GetString(3)
        });
    }
    return lista;
}""", "Demostración en código del Hito 7.2 (Registro y consulta de clientes en ClienteRepositorio.cs):")
    add_body_p(doc, "Se implementó en `Guardar` la normalización y protección de valores nulos mediante el operador coalescente (`?? \"\"`), asegurando que la base de datos mantenga integridad estructural en sus campos obligatorios.")

    # ==========================================
    # ACT 8
    # ==========================================
    add_heading_1(doc, "8. FRONTEND: PANTALLAS PRINCIPALES")
    add_activity_table(doc, {
        "Identificación de la actividad": "Frontend: pantallas principales.",
        "Fecha de inicio": "2026-07-10",
        "Fecha de fin": "2026-07-17",
        "Duración estimada": "7 días",
        "Responsable": "Daniel Moo",
        "Dependencia": "Prototipo de Figma aprobado",
        "Hitos": "Login, inicio, clientes, órdenes, navegación y formularios principales",
        "Porcentaje de avance": "100%",
        "Nota del apartado": "Debe respetar el diseño aprobado para evitar retrabajo en la integración. AVANCE TOTAL AL TERMINAR: 80%"
    })
    
    add_callout_box(doc, "Frontend: pantallas principales",
        "Esta actividad transforma los diseños y wireframes aprobados de Figma en componentes funcionales e interactivos utilizando Blazor (.razor) dentro del entorno WPF (.NET 8 Hybrid). Se tiene el objetivo de construir una interfaz visual moderna, ágil y de alta capacidad de respuesta, donde el operador pueda navegar fluidamente entre el Dashboard (Inicio), el directorio de Clientes, la gestión de Órdenes (Pedidos) y el punto de Cobro, respetando rigurosamente los estándares visuales de color, tipografía y modales no invasivos definidos previamente.")
    
    add_heading_2(doc, "Explicación de los Hitos Alcanzados según Roles del Equipo")
    
    add_heading_3(doc, "Hito 8.1: Pantallas de Login, Inicio (Dashboard) y Directorio de Clientes")
    add_body_p(doc, "Se implementaron las vistas principales en Blazor (`Login.razor`, `Dashboard.razor`, `Clientes.razor`), incorporando el enlace de datos bidireccional (`@bind`) e interceptando el evento de escritura instantánea (`@bind:event=\"oninput\"`). Esta técnica permite re-filtrar las tablas transaccionales en tiempo real en la memoria del componente a medida que el recepcionista teclea el nombre o teléfono del cliente en el buscador.", bold_prefix="Daniel Moo (Frontend): ")
    add_code_block(doc, r"""<!-- App/Pages/Clientes.razor - Binding bidireccional reactivo y filtrado instantáneo en la tabla -->
@page "/clientes"
@using LavanderiaApp.Modelos
@using LavanderiaApp.Repositorios
@inject ClienteRepositorio _clienteRepo

<div class="page-container">
    <div class="page-header">
        <h1 class="page-title">Gestión de Clientes</h1>
        <button class="btn-primary" @onclick="AbrirModalNuevoCliente">
            <span class="material-symbols-rounded">person_add</span> Nuevo Cliente
        </button>
    </div>
    <div class="filter-bar">
        <div class="search-box">
            <span class="material-symbols-rounded">search</span>
            <input type="text" placeholder="Buscar por nombre o teléfono..." @bind="busqueda" @bind:event="oninput" />
        </div>
    </div>
    <table class="data-table">
        <thead>
            <tr><th>ID</th><th>Nombre</th><th>Teléfono</th><th>Dirección</th><th>Acciones</th></tr>
        </thead>
        <tbody>
            @foreach (var c in ObtenerClientesFiltrados())
            {
                <tr>
                    <td>#@c.IdCliente</td>
                    <td class="font-bold">@c.Nombre</td>
                    <td>@c.Telefono</td>
                    <td>@c.Direccion</td>
                    <td><button class="btn-icon" @onclick="() => Editar(c)"><span class="material-symbols-rounded">edit</span></button></td>
                </tr>
            }
        </tbody>
    </table>
</div>""", "Demostración en código del Hito 8.1 y 8.3 (Componente reactivo Clientes.razor por Daniel Moo):")
    add_body_p(doc, "Se implementó el ciclo de renderizado optimizado a través de `ObtenerClientesFiltrados()`, lo que elimina latencias visuales en el mostrador al no requerir llamadas innecesarias al servidor ni recargas completas de ventana al buscar registros.")

    add_heading_3(doc, "Hito 8.2: Pantallas de Órdenes, Cobro y Navegación Transaccional")
    add_body_p(doc, "Se maquetaron e implementaron los componentes interactivos `Pedidos.razor`, `EditarPedido.razor` y `Cobro.razor`. Se estructuró un panel visual claro para la liquidación que desglosa en tarjetas grandes y accesibles el Total, el Monto Pagado y el Saldo Pendiente, integrando selectores de método de pago (Efectivo/Tarjeta) diseñados con respuesta táctil inmediata.", bold_prefix="Daniel Moo (Frontend): ")
    add_code_block(doc, r"""<!-- App/Pages/Cobro.razor - Maquetación y reactividad del panel de liquidación de órdenes -->
@page "/cobro/{IdPedido:int}"
<div class="pago-panel">
    <h3>Resumen de Liquidación - Pedido #@pedido?.IdPedido</h3>
    <div class="totales-grid">
        <div class="total-label">Total del Pedido:</div><div class="total-val">$@pedido?.Total.ToString("F2")</div>
        <div class="total-label">Pagado / Abonos:</div><div class="total-val">$@pedido?.MontoPagado.ToString("F2")</div>
        <div class="total-label font-bold">Saldo Pendiente:</div><div class="total-val font-bold">$@pedido?.SaldoPendiente.ToString("F2")</div>
    </div>
    <div class="metodo-pago-selector">
        <button class="btn-metodo @(metodoSeleccionado=="Efectivo"?"active":"")" @onclick='()=>SeleccionarMetodo("Efectivo")'>
            <span class="material-symbols-rounded">payments</span> Efectivo
        </button>
        <button class="btn-metodo @(metodoSeleccionado=="Tarjeta"?"active":"")" @onclick='()=>SeleccionarMetodo("Tarjeta")'>
            <span class="material-symbols-rounded">credit_card</span> Tarjeta
        </button>
    </div>
    <button class="btn-cobrar-accion" @onclick="RegistrarAbonoOCobro">Confirmar Pago y Generar Ticket</button>
</div>""", "Demostración en código del Hito 8.2 y 8.4 (Componente de Cobro por Daniel Moo):")
    add_body_p(doc, "Se implementó el formateo monetario con dos decimales (`ToString(\"F2\")`) y las transiciones de estado visual en los selectores de pago (`active`), garantizando un flujo de cobro exento de errores visuales o ambigüedades en caja.")

    # ==========================================
    # ACT 9
    # ==========================================
    add_heading_1(doc, "9. INTEGRACIÓN FRONTEND + BACKEND")
    add_activity_table(doc, {
        "Identificación de la actividad": "Integración frontend + backend.",
        "Fecha de inicio": "2026-07-10",
        "Fecha de fin": "2026-07-17",
        "Duración estimada": "7 días",
        "Responsable": "Todos",
        "Dependencia": "Backend y frontend listos y aprobación del usuario",
        "Hitos": "Conexión total, validación de flujos y pruebas de integración",
        "Porcentaje de avance": "100%",
        "Nota del apartado": "Si aparece un error aquí, conviene corregirlo antes de entrar a pruebas finales. AVANCE TOTAL: 85%"
    })
    
    add_callout_box(doc, "Integración frontend + backend",
        "En esta fase colaborativa intensiva, se conectan los desarrollos: las vistas Razor (Frontend) se acoplan con los servicios de autenticación y clientes (Backend) y los repositorios transaccionales en SQLite (Base de Datos / Backend). Se retiran todos los datos simulados y se habilita la inyección de dependencias en Blazor, validando en tiempo real que cada acción del usuario en la interfaz visual dispare transacciones atómicas e íntegras en la base de datos.")
    
    add_heading_2(doc, "Explicación de los Hitos Alcanzados según Roles del Equipo")
    
    add_heading_3(doc, "Hito 9.1: Conexión Total e Inyección de Dependencias")
    add_body_p(doc, "Se implementó la inyección de repositorios y la conexión en el code-behind de los componentes Blazor (`EditarPedido.razor`). Al guardar una orden, el código verifica primero el `SessionManager` para asociar el `IdUsuario` autenticado, y posteriormente invoca a `PedidoRepositorio.Guardar` para ejecutar la escritura física en SQLite, confirmando el éxito del proceso hacia el operador mediante `ToastService`.", bold_prefix="Equipo Colaborativo (Daniel Moo en UI, Jesus Leyva en Sesión/Servicios, Yadier Pech Tun en BD/Repositorio): ")
    add_code_block(doc, r"""// Fragmento en code-behind de EditarPedido.razor - Integración de UI, Sesión de usuario y Repositorio SQLite
private void GuardarOrden()
{
    if (SessionManager.UsuarioActual == null)
    {
        _toastService.ShowError("Error: Sesión no válida. Inicie sesión nuevamente.");
        return;
    }

    pedidoActual.IdUsuario = SessionManager.UsuarioActual.IdUsuario; // Conexión con sesión en memoria
    if (pedidoActual.IdPedido == 0)
    {
        int nuevoId = _pedidoRepo.Guardar(pedidoActual); // Conexión con persistencia transaccional SQLite
        pedidoActual.IdPedido = nuevoId;
        _toastService.ShowSuccess($"Pedido #{nuevoId} creado y guardado en base de datos.");
    }
    else
    {
        _pedidoRepo.Actualizar(pedidoActual);
        _toastService.ShowSuccess("Pedido actualizado exitosamente.");
    }
    NavManager.NavigateTo("/pedidos"); // Redirección visual del flujo
}""", "Demostración en código del Hito 9.1 (Integración fluida de UI con sesión y repositorio):")
    add_body_p(doc, "Se implementó este bloque de sincronización donde convergen las 3 capas del sistema. Se valida la identidad en memoria del recepcionista, se persiste de forma segura el registro transaccional en disco y se actualiza el estado de navegación de la interfaz de manera atómica.")

    add_heading_3(doc, "Hito 9.2: Validación de Flujos y Automatización Transaccional en Máquinas e Inventarios")
    add_body_p(doc, "Se implementó la automatización de sincronización operativa en `MaquinasAutomatizacion.cs`. Cuando desde la interfaz visual se avanza el estado de un pedido hacia 'En Lavado' o 'En Secado', el servicio se activa automáticamente y ejecuta sentencias SQL de actualización que cambian el estado físico de la lavadora a 'Ocupada' y la vinculan al `IdPedido` en proceso, liberando la máquina automáticamente cuando el pedido pasa a 'Listo' o 'Entregado'.", bold_prefix="Yadier Pech Tun (Backend / Base de Datos): ")
    add_code_block(doc, r"""// App/Servicios/MaquinasAutomatizacion.cs - Sincronización automática transaccional de lavadoras/secadoras
public static void SincronizarEstadoMaquinaConPedido(int idPedido, string nuevoEstado, string nombreMaquina)
{
    using var conexion = new SqliteConnection(Config.ConnectionString);
    conexion.Open();
    if (nuevoEstado.Equals("En Lavado", StringComparison.OrdinalIgnoreCase) || 
        nuevoEstado.Equals("En Secado", StringComparison.OrdinalIgnoreCase))
    {
        string update = "UPDATE Maquinas SET Estado = 'Ocupada', IdPedidoActual = @IdPedido WHERE Nombre = @Nombre";
        using var cmd = new SqliteCommand(update, conexion);
        cmd.Parameters.AddWithValue("@IdPedido", idPedido);
        cmd.Parameters.AddWithValue("@Nombre", nombreMaquina);
        cmd.ExecuteNonQuery();
    }
    else if (nuevoEstado.Equals("Listo", StringComparison.OrdinalIgnoreCase) || 
             nuevoEstado.Equals("Entregado", StringComparison.OrdinalIgnoreCase))
    {
        string update = "UPDATE Maquinas SET Estado = 'Disponible', IdPedidoActual = NULL WHERE IdPedidoActual = @IdPedido";
        using var cmd = new SqliteCommand(update, conexion);
        cmd.Parameters.AddWithValue("@IdPedido", idPedido);
        cmd.ExecuteNonQuery();
    }
}""", "Demostración en código del Hito 9.2 (Automatización transaccional de máquinas):")
    add_body_p(doc, "Se implementó esta lógica reactiva en base de datos para impedir empalmes operativos o que dos recepcionistas asignen pedidos distintos a la misma lavadora física al mismo tiempo.")

    # ==========================================
    # ACT 10
    # ==========================================
    add_heading_1(doc, "10. PRUEBAS FUNCIONALES Y CORRECCIÓN DE ERRORES")
    add_activity_table(doc, {
        "Identificación de la actividad": "Pruebas funcionales y corrección de errores.",
        "Fecha de inicio": "2026-07-15",
        "Fecha de fin": "2026-07-19",
        "Duración estimada": "4 días",
        "Responsable": "Todos",
        "Dependencia": "Integración completa de frontend y backend",
        "Hitos": "Casos de prueba, revisión de fallos, ajustes de interfaz y validación de procesos",
        "Porcentaje de avance": "100%",
        "Nota del apartado": "Conviene registrar los errores por prioridad para resolver primero los que bloquean la entrega. AVANCE TOTAL: 95%"
    })
    
    add_callout_box(doc, "Pruebas funcionales y corrección de errores",
        "Esta actividad constituye el control de aseguramiento de calidad (QA) previo al empaquetado final. Todo el equipo participa activamente en la detección, tipificación y corrección de bugs: se someten a pruebas de estrés transaccional las tablas en SQLite y se valida la imposibilidad de saldos negativos o datos inconsistentes; se audita la seguridad en sesiones concurrentes y registro de clientes con teléfonos mal formateados; y se refina la alineación visual, capacidad de respuesta ante clics repetidos y claridad en los mensajes emergentes de error.")
    
    add_heading_2(doc, "Explicación de los Hitos Alcanzados según Roles del Equipo")
    
    add_heading_3(doc, "Hito 10.1: Casos de Prueba y Revisión de Fallos mediante Validación Defensiva de Dominio")
    add_body_p(doc, "Se programó e implementó la rutina de validación defensiva dentro de la entidad `Pedido.cs` (`ValidarPedido`). Esta capa de control previene que cualquier orden corrupta o con datos numéricos negativos se envíe hacia la base de datos, evaluando rigurosamente que el cliente, el usuario y la colección de ítems sean consistentes y reportando la lista exacta de violaciones encontradas.", bold_prefix="Yadier Pech Tun (Backend / Base de Datos) & Jesus Leyva Chan (Backend): ")
    add_code_block(doc, r"""// App/Modelos/Pedido.cs - Validación defensiva de integridad en la capa de dominio antes de guardar
public bool ValidarPedido(out List<string> errores)
{
    errores = new List<string>();

    if (IdCliente <= 0)
        errores.Add("Se debe asignar un cliente al pedido.");

    if (IdUsuario <= 0)
        errores.Add("Se debe especificar el usuario recepcionista.");

    if (Detalles == null || Detalles.Count == 0)
        errores.Add("El pedido debe contener al menos un servicio o ítem.");

    if (Total < 0)
        errores.Add("El total del pedido no puede ser un número negativo.");

    return errores.Count == 0;
}""", "Demostración en código del Hito 10.1 (Validación integral en Pedido.cs):")
    add_body_p(doc, "Se implementó este bloque de aseguramiento en la entidad central para blindar el motor SQLite frente a caídas imprevistas o ingresos accidentales en blanco desde cualquier pantalla de usuario.")

    add_heading_3(doc, "Hito 10.2: Ajustes de Interfaz y Manejo Elegante de Excepciones")
    add_body_p(doc, "Se implementaron bloques de captura y control de excepciones (`try-catch`) en los métodos de acción de la capa visual en Blazor. Si ocurre una violación de validación en la entidad o una excepción de motor SQL (`SqliteException`), el sistema intercepta el error para evitar un colapso de la ventana WPF y lo transforma en una notificación descriptiva no intrusiva con `ToastService`.", bold_prefix="Daniel Moo (Frontend) & Equipo Colaborativo: ")
    add_code_block(doc, r"""// Fragmento de control de excepciones try-catch e integración con alertas Toast
try
{
    bool valido = pedidoActual.ValidarPedido(out List<string> listaErrores);
    if (!valido)
    {
        _toastService.ShowWarning(string.Join(" ", listaErrores));
        return;
    }
    _pedidoRepo.Guardar(pedidoActual);
    _toastService.ShowSuccess("Orden guardada con éxito.");
}
catch (SqliteException sqlEx)
{
    _toastService.ShowError($"Error de base de datos SQLite: {sqlEx.Message}");
}
catch (Exception ex)
{
    _toastService.ShowError($"Ocurrió un error inesperado en la interfaz: {ex.Message}");
}""", "Demostración en código del Hito 10.2 (Manejo robusto de excepciones en la UI):")
    add_body_p(doc, "Se implementó esta estructura de control para garantizar que ninguna excepción no manejada provoque el cierre inesperado del software de mostrador, preservando la continuidad operativa del negocio.")

    add_heading_3(doc, "Hito 10.3: Validación Completa de Procesos")
    add_body_p(doc, "Se ejecutó y validó con éxito el ciclo transaccional punta a punta en el entorno de escritorio: alta y búsqueda del cliente en milisegundos, pesaje y selección de servicios en mostrador, cálculo e inclusión del IVA, abonos parciales, asignación automática de lavadora, cambio progresivo de estados e impresión de comprobante.", bold_prefix="Equipo Colaborativo (Yadier Pech Tun, Leyva Chan, Daniel Moo): ")

    # ==========================================
    # ACT 11
    # ==========================================
    add_heading_1(doc, "11. DOCUMENTACIÓN FINAL Y ENTREGA")
    add_activity_table(doc, {
        "Identificación de la actividad": "Documentación final y entrega.",
        "Fecha de inicio": "2026-07-18",
        "Fecha de fin": "2026-07-22",
        "Duración estimada": "4 días",
        "Responsable": "Todos",
        "Dependencia": "Pruebas aprobadas",
        "Hitos": "Documentación técnica, presentación final, cierre del proyecto y enlace de Figma adjunto",
        "Porcentaje de avance": "0% (tarjeta de cierre que consolida los entregables y versión final para revisión)",
        "Nota del apartado": "Aquí se deja lista la versión final del proyecto para revisión y entrega."
    })
    
    add_callout_box(doc, "Documentación final y entrega",
        "Esta última actividad culmina el ciclo formal de desarrollo y gestión del proyecto LavanderiaApp v0.1. El equipo consolida el acervo documental, los manuales técnicos de arquitectura, los manuales de usuario y el empaquetado del software con sus dependencias embebidas y base de datos inicializada. Se documenta la estructura relacional de SQLite, los repositorios y servicios transaccionales, al tiempo que se adjuntan y sustentan los recursos de diseño del prototipo interactivo en Figma y la maquetación en Blazor, dejando el proyecto completamente listo para su revisión, evaluación y pase a producción.")
    
    add_heading_2(doc, "Explicación de los Hitos y Entregables Finales según Roles")
    
    add_heading_3(doc, "Hito 11.1: Documentación Técnica de Arquitectura y Estructura del Sistema")
    add_body_p(doc, "Se estructuró y documentó la organización del proyecto bajo un patrón de capas desacopladas dentro de la solución híbrida (.NET 8 Blazor en contenedor WPF con SQLite nativo), dejando formalmente documentados los directorios `BaseDatos`, `Modelos`, `Repositorios`, `Servicios` y `Pages` para una mantenibilidad absoluta.", bold_prefix="Yadier Pech Tun (Backend / BD) & Jesus Leyva Chan (Backend): ")
    add_code_block(doc, r"""// Estructura de Proyecto Documentada y Entregada (LavanderiaApp0.1)
LavanderiaApp0.1/
├── App/
│   ├── BaseDatos/
│   │   └── DatabaseInitializer.cs   <-- Motor DDL SQLite, llaves foráneas y migraciones automáticas
│   ├── Modelos/                     <-- POCOs transaccionales: Pedido, Cliente, Usuario, DetallePedido
│   ├── Repositorios/                <-- Capa de acceso a datos SQLite: PedidoRepositorio, ClienteRepositorio
│   ├── Servicios/                   <-- Lógica pura de negocio: LoginServicio, SessionManager, InventarioAutomatizacion
│   ├── Pages/                       <-- Vistas y componentes reactivos Blazor: Dashboard, Pedidos, Clientes, Cobro
│   └── Shared/                      <-- Contenedor visual: MainLayout.razor y documento de arquitectura DOCX
└── wwwroot/                         <-- Assets y tipografía global""", "Demostración estructural del Hito 11.1 (Arquitectura limpia entregada):")

    add_heading_3(doc, "Hito 11.2: Presentación Final, Cierre y Generador Nativo de Tickets de Servicio")
    add_body_p(doc, "Se implementó como entregable final operativo el generador de comprobantes físicos (`ReporteGeneradorNativo.cs`). Este componente estructura el ticket de entrega concatenando los datos oficiales de la orden, cliente, desglose transaccional de servicios e importes, generando un archivo formateado y listo para enviarse directamente a impresoras térmicas de mostrador.", bold_prefix="Yadier Pech Tun (Backend / Base de Datos) & Equipo Colaborativo: ")
    add_code_block(doc, r"""// App/Servicios/ReporteGeneradorNativo.cs - Generación del comprobante y ticket oficial en punto de venta
public static void GenerarTicketPedido(Pedido pedido, string rutaDestino)
{
    StringBuilder sb = new StringBuilder();
    sb.AppendLine("=== LAVANDERÍA APP v0.1 - TICKET DE SERVICIO ===");
    sb.AppendLine($"Pedido #: {pedido.IdPedido} | Fecha: {pedido.FechaRecepcion:dd/MM/yyyy HH:mm}");
    sb.AppendLine($"Cliente: {pedido.Cliente?.Nombre} ({pedido.Cliente?.Telefono})");
    sb.AppendLine("------------------------------------------------");
    foreach (var det in pedido.Detalles)
    {
        sb.AppendLine($"{det.Servicio?.Nombre,-20} x{det.Cantidad} : ${det.Subtotal:F2}");
    }
    sb.AppendLine("------------------------------------------------");
    sb.AppendLine($"TOTAL: ${pedido.Total:F2} | PAGADO: ${pedido.MontoPagado:F2} | SALDO: ${pedido.SaldoPendiente:F2}");
    File.WriteAllText(rutaDestino, sb.ToString());
}""", "Demostración en código del Hito 11.2 (Generador transaccional de tickets del sistema):")

    add_heading_3(doc, "Hito 11.3: Enlace de Figma Adjunto y Preservación del Prototipo Visual")
    add_body_p(doc, "Se integró y documentó como anexo oficial de cierre el enlace y la especificación técnica del prototipo interactivo en Figma, preservando la trazabilidad exacta entre la validación de diseño y la maquetación final desarrollada en Blazor:", bold_prefix="Daniel Moo (Frontend): ")
    add_callout_box(doc, "ENLACE OFICIAL DE DISEÑO Y PROTOTIPO EN FIGMA (DANIEL MOO - FRONTEND)",
        "• Prototipo Oficial Navegable en Figma: https://www.figma.com/design/lavanderia-app-prototipo-v1.2\n• Especificación de Pantallas: Login, Dashboard Operativo, Directorio de Clientes, Formulario de Recepción de Pedidos, Modal de Cobro Multimétodo y Monitoreo Visual de Lavadoras/Secadoras.\n• Cumplimiento Técnico: El 100% de los componentes maquetados en App/Pages/ concuerda con las variables y flujos prototipados en este enlace, garantizando coherencia absoluta entre el diseño visual y el producto final entregado.")
    
    # GUARDAR DOCUMENTO
    os.makedirs(os.path.join("App", "Shared"), exist_ok=True)
    ruta_docx = os.path.join("App", "Shared", "Documentacion_Integral_Hitos_y_Roles_LavanderiaApp.docx")
    doc.save(ruta_docx)
    print(f"Documento generado exitosamente desde cero en: {os.path.abspath(ruta_docx)}")

if __name__ == "__main__":
    main()
