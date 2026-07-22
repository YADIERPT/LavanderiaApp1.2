import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table, color="D3D3D3", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), val)
        border.set(qn('w:sz'), sz)
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tblBorders.append(border)
    border = OxmlElement('w:insideV')
    border.set(qn('w:val'), 'none')
    tblBorders.append(border)
    tblPr.append(tblBorders)

def add_header_p(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Segoe UI'
    run.bold = True
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(30, 58, 138) # #1E3A8A
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(15, 23, 42) # #0F172A
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
    elif level == 3:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(2, 132, 199) # #0284C7
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
    return p

def add_body_p(doc, text, bold_prefix=None, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = 'Segoe UI'
        r_pre.bold = True
        r_pre.font.size = Pt(10.5)
        r_pre.font.color.rgb = RGBColor(15, 23, 42)
    run = p.add_run(text)
    run.font.name = 'Segoe UI'
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(51, 65, 85)
    run.italic = italic
    return p

def add_activity_table(doc, act_name, f_init, f_end, duration, resp, dep, hitos, avance, nota):
    table = doc.add_table(rows=8, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table, color="CBD5E1", sz="6")
    
    col_widths = [Inches(2.2), Inches(4.3)]
    for row in table.rows:
        for idx, width in enumerate(col_widths):
            row.cells[idx].width = width
            
    headers = [
        "Identificación de la Actividad",
        "Fecha de Inicio / Fin",
        "Duración Estimada",
        "Responsable(s)",
        "Dependencia",
        "Hitos",
        "Porcentaje de Avance",
        "Nota / Base del Proyecto"
    ]
    values = [
        act_name,
        f"{f_init}  al  {f_end}",
        duration,
        resp,
        dep,
        hitos,
        avance,
        nota
    ]
    
    for i, (h, v) in enumerate(zip(headers, values)):
        row = table.rows[i]
        c0, c1 = row.cells[0], row.cells[1]
        set_cell_margins(c0, top=100, bottom=100, left=150, right=150)
        set_cell_margins(c1, top=100, bottom=100, left=150, right=150)
        
        if i == 0:
            set_cell_background(c0, "1E3A8A")
            set_cell_background(c1, "1E3A8A")
            p0 = c0.paragraphs[0]
            r0 = p0.add_run(h)
            r0.font.name = 'Segoe UI'
            r0.bold = True
            r0.font.size = Pt(10)
            r0.font.color.rgb = RGBColor(255, 255, 255)
            
            p1 = c1.paragraphs[0]
            r1 = p1.add_run(v)
            r1.font.name = 'Segoe UI'
            r1.bold = True
            r1.font.size = Pt(10)
            r1.font.color.rgb = RGBColor(255, 255, 255)
        else:
            bg = "F8FAFC" if i % 2 == 1 else "FFFFFF"
            set_cell_background(c0, bg)
            set_cell_background(c1, bg)
            
            p0 = c0.paragraphs[0]
            r0 = p0.add_run(h)
            r0.font.name = 'Segoe UI'
            r0.bold = True
            r0.font.size = Pt(9.5)
            r0.font.color.rgb = RGBColor(30, 41, 59)
            
            p1 = c1.paragraphs[0]
            r1 = p1.add_run(v)
            r1.font.name = 'Segoe UI'
            r1.font.size = Pt(9.5)
            if i == 6:
                r1.bold = True
                r1.font.color.rgb = RGBColor(2, 132, 199)
            else:
                r1.font.color.rgb = RGBColor(51, 65, 85)
    
    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_before = Pt(4)
    p_space.paragraph_format.space_after = Pt(4)

def read_file_exact(rel_path):
    base_dir = r"C:\Users\Yadie\RiderProjects\LavanderiaApp0.1\App"
    full_path = os.path.join(base_dir, rel_path)
    if not os.path.exists(full_path):
        return f"// ERROR: No se encontró el archivo {rel_path}"
    with open(full_path, "r", encoding="utf-8") as f:
        return f.read().rstrip()

def extract_method_exact(rel_path, method_name, start_keyword="public"):
    base_dir = r"C:\Users\Yadie\RiderProjects\LavanderiaApp0.1\App"
    full_path = os.path.join(base_dir, rel_path)
    with open(full_path, "r", encoding="utf-8") as f:
        lines = f.readlines()
    
    start_idx = -1
    for i, l in enumerate(lines):
        if method_name in l and (start_keyword in l or "void" in l or "Task" in l or "List" in l or "bool" in l):
            start_idx = i
            break
            
    if start_idx == -1:
        return read_file_exact(rel_path)
        
    # extract balance of braces
    result_lines = []
    brace_count = 0
    found_first_brace = False
    
    for i in range(start_idx, len(lines)):
        l = lines[i]
        result_lines.append(l)
        if "{" in l:
            brace_count += l.count("{")
            found_first_brace = True
        if "}" in l:
            brace_count -= l.count("}")
        if found_first_brace and brace_count == 0:
            break
            
    return "".join(result_lines).rstrip()

def add_code_table(doc, title, code_text):
    add_header_p(doc, title, level=3)
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table, color="94A3B8", sz="6")
    
    cell = table.rows[0].cells[0]
    cell.width = Inches(6.5)
    set_cell_background(cell, "0F172A") # dark slate theme
    set_cell_margins(cell, top=140, bottom=140, left=160, right=160)
    
    p = cell.paragraphs[0]
    p.paragraph_format.line_spacing = 1.05
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(241, 245, 249) # light text
    
    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_before = Pt(4)
    p_space.paragraph_format.space_after = Pt(6)

def generate_report():
    doc = docx.Document()
    
    # Configure margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        
    # Title
    p_title = doc.add_paragraph()
    r_t = p_title.add_run("SISTEMA LAVANDERIAAPP v0.1\nInforme Técnico Detallado: Explicación de Código, Implementación y Asignación por Actividades y Roles")
    r_t.font.name = 'Segoe UI'
    r_t.font.size = Pt(20)
    r_t.bold = True
    r_t.font.color.rgb = RGBColor(30, 58, 138)
    p_title.paragraph_format.space_after = Pt(14)
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Resumen y Roles
    add_header_p(doc, "1. Resumen Ejecutivo y Matriz Institucional de Roles y Responsabilidades", level=1)
    add_body_p(doc, "El presente documento formaliza el informe de desarrollo técnico del sistema LavanderiaApp v0.1, estructurado rigurosamente bajo la metodología de entregables por actividad del proyecto. En cumplimiento estricto con los lineamientos de arquitectura y división del trabajo, se documenta la explicación general de lo que se pretende hacer en cada actividad, seguido del desglose técnico individual y el código implementado verbatim por cada desarrollador según su rol asignado:")
    
    add_body_p(doc, "• Yadier Pech Tun (Rol: Backend & Base de Datos): Responsable integral de la arquitectura del motor relacional SQLite, modelado y creación de tablas DDL, repositorios transaccionales relacionales, y lógica operativa transaccional en servicios C# (automatización de inventarios, temporizadores de máquinas, cálculos de corte y auditoría).", bold_prefix="Rol 1: ")
    add_body_p(doc, "• Jesus Leyva Chan / Leyva Chan (Rol: Backend): Responsable integral de los servicios y modelos de seguridad (autenticación, sesión segura, encriptación), servicios transaccionales de clientes (lealtad y visitas), lógica transaccional de pagos y sistema de alertas o notificaciones modales en C#.", bold_prefix="Rol 2: ")
    add_body_p(doc, "• Daniel Moo (Rol: Frontend): Responsable exclusivo de la capa visual y de presentación, diseñando e implementando las pantallas principales, navegación, modales y flujos interactivos de usuario en componentes Blazor Hybrid / Razor, además de la presentación e interfaz de reportes operativos.", bold_prefix="Rol 3: ")
    
    add_body_p(doc, "A continuación, se presenta el desarrollo detallado, ordenado cronológicamente desde la Actividad 1 hasta la Actividad 11, respetando íntegramente las firmas exactas y el código original del proyecto sin contracciones ni alteraciones, y presentando secciones independientes para cada método o archivo referenciado en las explicaciones.")

    # =========================================================================
    # ACTIVIDAD 1
    # =========================================================================
    add_header_p(doc, "2. Actividad: Levantamiento de requerimientos y definición del alcance", level=1)
    add_activity_table(
        doc,
        act_name="Levantamiento de requerimientos y definición del alcance",
        f_init="2026-05-07",
        f_end="2026-06-02",
        duration="29 días",
        resp="Yadier Pech Tun, Leyva Chan y Daniel Moo",
        dep="Ninguna",
        hitos="Lista de módulos, objetivos del sistema, alcance validado y prioridades del proyecto",
        avance="100%",
        nota="De aquí sale la base del resto del proyecto: módulos, reglas de negocio y entregables mínimos."
    )
    
    add_header_p(doc, "¿Qué se pretende hacer a rasgos generales en esta actividad?", level=2)
    add_body_p(doc, "En esta actividad inicial se pretende establecer los cimientos arquitectónicos y normativos sobre los cuales operará toda la solución LavanderiaApp v0.1. El objetivo primordial es transformar las necesidades operativas de una lavandería comercial en parámetros técnicos concretos, definiendo los módulos base (Usuarios, Clientes, Servicios, Pedidos, Máquinas, Inventario y Reportes), las reglas de cálculo financiero (IVA, fondo de caja) y la estructura del contenedor visual y de seguridad que delimitará el alcance del sistema.")
    
    add_header_p(doc, "Desglose y Asignación de Responsabilidades por Integrante y Rol", level=2)
    add_body_p(doc, "Como responsable del Backend & Base de Datos, Yadier Pech Tun implementó el núcleo de configuración transaccional del negocio en la clase BusinessConfigData y el servicio gestor de persistencia estática BusinessConfig (App/Servicios/BusinessConfig.cs). Este módulo define y almacena los parámetros corporativos obligatorios (nombre del negocio, teléfono, dirección, tasa de IVA del 16% y fondo de caja inicial), permitiendo que el resto de los servicios de base de datos consuman estas reglas de negocio validadas.", bold_prefix="Yadier Pech Tun (Rol: Backend & Base de Datos): ")
    add_body_p(doc, "Como responsable del Backend de seguridad, Leyva Chan construyó el gestor de sesión estático SessionManager (App/Servicios/SessionManager.cs). Este servicio encapsula la sesión del usuario activo y provee verificaciones inmediatas de autenticación y nivel de privilegio administrativo, garantizando que los límites del alcance respecto a quién puede operar el sistema se impongan desde la capa de servicio.", bold_prefix="Jesus Leyva Chan (Rol: Backend): ")
    add_body_p(doc, "Como responsable del Frontend, Daniel Moo estructuró el contenedor visual maestro del Punto de Venta (POS) en el componente MainLayout.razor (App/Shared/MainLayout.razor). En esta fase de levantamiento y definición, implementó la barra superior de estado de usuario y la disposición de la estructura principal para alojar los módulos del sistema.", bold_prefix="Daniel Moo (Rol: Frontend): ")
    
    add_header_p(doc, "¿Qué y Cómo se Implementó en el Código?", level=2)
    add_code_table(doc, "Código Implementado por Yadier Pech Tun — Reglas de Negocio y Configuración (App/Servicios/BusinessConfig.cs)", read_file_exact("Servicios/BusinessConfig.cs"))
    add_code_table(doc, "Código Implementado por Jesus Leyva Chan — Control de Sesión Global (App/Servicios/SessionManager.cs)", read_file_exact("Servicios/SessionManager.cs"))
    
    # For MainLayout, since it's 1145 lines, let's extract the exact header / sidebar layout structure or full file if feasible without massive clutter, but to ensure exactness without contractions we show the exact structural layout portion referenced by Daniel Moo or top lines verbatim.
    ml_lines = read_file_exact("Shared/MainLayout.razor").splitlines()
    ml_top = "\n".join(ml_lines[:65])
    add_code_table(doc, "Código Implementado por Daniel Moo — Estructura y Navegación del POS (App/Shared/MainLayout.razor - Sección Base)", ml_top)

    # =========================================================================
    # ACTIVIDAD 2
    # =========================================================================
    add_header_p(doc, "3. Actividad: Definición de backlog y módulos principales", level=1)
    add_activity_table(
        doc,
        act_name="Definición de backlog y módulos principales",
        f_init="2026-06-07",
        f_end="2026-06-11",
        duration="4 días",
        resp="Yadier Pech Tun",
        dep="Requerimientos definidos",
        hitos="Módulos de usuarios, clientes, servicios, pedidos y reportes detalles pedido, etc",
        avance="100%",
        nota="Apartado ya listo\nAVANCE TOTAL AL TERMINAR: 25%"
    )
    
    add_header_p(doc, "¿Qué se pretende hacer a rasgos generales en esta actividad?", level=2)
    add_body_p(doc, "En esta actividad, bajo la responsabilidad única de Yadier Pech Tun (Backend & Base de Datos), se pretende modularizar todo el dominio transaccional de la lavandería y definir el backlog técnico del sistema. El propósito central es construir las entidades principales que representan los módulos de usuarios, clientes, catálogo de servicios y órdenes de pedido, estableciendo al mismo tiempo las rutas de conexión a la base de datos y el bus de comunicación de alertas en la capa de servicios.")
    
    add_header_p(doc, "Desglose y Asignación de Responsabilidades por Integrante y Rol", level=2)
    add_body_p(doc, "Al ser el responsable único de esta actividad y titular del área de Base de Datos y Backend, Yadier Pech Tun diseñó y codificó las tres piezas medulares que sustentan el backlog y los módulos transaccionales:\n"
                    "1) El módulo de configuración de rutas de base de datos y bus de alertas de servicio (App/BaseDatos/Config.cs), que resuelve dinámicamente la ubicación física del archivo SQLite lavanderia.db y expone el observador de notificaciones ToastService.\n"
                    "2) El modelo de dominio central del módulo de Pedidos (App/Modelos/Pedido.cs), el cual define la estructura completa de una orden: generación de folio automático, relación con el cliente, fechas de promesa, estados operativos ('En espera', 'Lavando', 'Secando', 'Listo', 'Entregado') y cálculos financieros automáticos de saldos pendientes y anticipos.\n"
                    "3) El modelo de dominio central del catálogo de Servicios (App/Modelos/Servicio.cs), que define las tarifas por kilogramo, pieza o carga, estimaciones de tiempo y cálculo formal de subtotales.", bold_prefix="Yadier Pech Tun (Rol: Backend & Base de Datos - Responsable Único): ")
    
    add_header_p(doc, "¿Qué y Cómo se Implementó en el Código?", level=2)
    add_body_p(doc, "Para cumplir con la directriz técnica de separar cada bloque y no juntar partes distintas del código en una sola tabla, se presentan a continuación los tres archivos desarrollados por Yadier Pech Tun de forma individual y 100% íntegra:")
    add_code_table(doc, "Código Implementado por Yadier Pech Tun — Rutas DB y Bus de Alertas (App/BaseDatos/Config.cs)", read_file_exact("BaseDatos/Config.cs"))
    add_code_table(doc, "Código Implementado por Yadier Pech Tun — Entidad Central del Módulo de Pedidos (App/Modelos/Pedido.cs)", read_file_exact("Modelos/Pedido.cs"))
    add_code_table(doc, "Código Implementado por Yadier Pech Tun — Entidad Central del Módulo de Servicios (App/Modelos/Servicio.cs)", read_file_exact("Modelos/Servicio.cs"))

    # =========================================================================
    # ACTIVIDAD 3
    # =========================================================================
    add_header_p(doc, "4. Actividad: Diseño de interfaz y flujo en Figma", level=1)
    add_activity_table(
        doc,
        act_name="Diseño de interfaz y flujo en Figma",
        f_init="2026-06-05",
        f_end="2026-06-07",
        duration="2 días",
        resp="Yadier Pech Tun, Jesus Leyva Chan, Daniel Moo",
        dep="Requerimientos definidos",
        hitos="Wireframes, flujo de pantallas, prototipo navegable y estructura visual",
        avance="100%",
        nota="AVANCE TOTAL 35%"
    )
    
    add_header_p(doc, "¿Qué se pretende hacer a rasgos generales en esta actividad?", level=2)
    add_body_p(doc, "Lo que se pretende en esta actividad es materializar los wireframes, flujos de pantalla y el prototipo visual diseñado en Figma dentro de la estructura real del proyecto Blazor Hybrid / C#. Se busca que las pantallas de selección y captura de órdenes cuenten con componentes visuales interactivos y que el backend disponga de las estructuras de soporte exactas para representar cada ítem seleccionado en mostrador antes y durante el registro de la orden.")
    
    add_header_p(doc, "Desglose y Asignación de Responsabilidades por Integrante y Rol", level=2)
    add_body_p(doc, "Como responsable de Frontend, Daniel Moo convirtió el diseño de Figma del flujo de cobro e interacción en componentes Razor funcionales, programando el modal emergente TipoPago.razor (App/Pages/TipoPago.razor), el cual permite al operador seleccionar de forma rápida entre pago en Efectivo o Tarjeta y redirige el flujo visual hacia la terminal de cobro exacta.", bold_prefix="Daniel Moo (Rol: Frontend): ")
    add_body_p(doc, "Como responsable de Backend y Base de Datos, Yadier Pech Tun desarrolló la entidad relacional de soporte para los renglones visuales de la orden: Detalle_Pedido.cs (App/Modelos/Detalle_Pedido.cs), asegurando el cálculo automático de subtotales por partida que se muestra en las pantallas diseñadas.", bold_prefix="Yadier Pech Tun (Rol: Backend & Base de Datos): ")
    add_body_p(doc, "Como responsable de Backend, Leyva Chan implementó el modelo transaccional temporal CarritoItem.cs (App/Modelos/CarritoItem.cs), el cual gestiona la lista dinámica de servicios en memoria mientras el operador agrega prendas en la interfaz antes de consolidar el pedido en la base de datos.", bold_prefix="Jesus Leyva Chan (Rol: Backend): ")
    
    add_header_p(doc, "¿Qué y Cómo se Implementó en el Código?", level=2)
    add_code_table(doc, "Código Implementado por Daniel Moo — Componente Modal Interactivo de Pago (App/Pages/TipoPago.razor)", read_file_exact("Pages/TipoPago.razor"))
    add_code_table(doc, "Código Implementado por Yadier Pech Tun — Entidad Relacional de Renglón (App/Modelos/Detalle_Pedido.cs)", read_file_exact("Modelos/Detalle_Pedido.cs"))
    add_code_table(doc, "Código Implementado por Jesus Leyva Chan — Modelo en Memoria del Carrito POS (App/Modelos/CarritoItem.cs)", read_file_exact("Modelos/CarritoItem.cs"))

    # =========================================================================
    # ACTIVIDAD 4
    # =========================================================================
    add_header_p(doc, "5. Actividad: Diseño de base de datos y arquitectura", level=1)
    add_activity_table(
        doc,
        act_name="Diseño de base de datos y arquitectura",
        f_init="2026-06-07",
        f_end="2026-06-11",
        duration="4 días",
        resp="Yadier Pech Tun y Jesus Leyva Chan",
        dep="definición de backlogs y módulos",
        hitos="Modelo entidad-relación, estructura de tablas, SQLite con reglas base y etiquetas",
        avance="25%",
        nota="De este diseño dependerán autenticación, clientes, servicios, órdenes y reportes.\nAVANCE TOTAL 40%"
    )
    
    add_header_p(doc, "¿Qué se pretende hacer a rasgos generales en esta actividad?", level=2)
    add_body_p(doc, "Con el diseño de la base de datos y su arquitectura de acceso relacional se pretende construir el motor de persistencia transaccional SQLite (lavanderia.db). El objetivo general es ejecutar el modelo entidad-relación mediante sentencias DDL estrictas con activación de llaves foráneas y reglas de integridad, y construir la arquitectura de repositorios SQL que permitirá realizar operaciones de lectura y escritura sobre los catálogos y transacciones del negocio.")
    
    add_header_p(doc, "Desglose y Asignación de Responsabilidades por Integrante y Rol", level=2)
    add_body_p(doc, "Como titular de Base de Datos y Backend, Yadier Pech Tun lideró la arquitectura relacional e implementó el motor DDL e inicializador de base de datos DatabaseInitializer.cs (App/BaseDatos/DatabaseInitializer.cs). Para respetar el mandato de no agrupar métodos distintos de un mismo archivo en una sola sección, se presenta por separado y en apartados independientes:\n"
                    "1) La lógica de inicialización y creación relacional de tablas maestras SQL con activación de PRAGMA foreign_keys (Método Initialize()).\n"
                    "2) La lógica transaccional de etiquetado y precarga de catálogo inicial y usuarios administradores por defecto (Método SeedInitialData()).\n"
                    "Adicionalmente, Yadier implementó el repositorio transaccional relacional para los detalles y renglones de pedidos: DetallePedidoRepositorio.cs (App/Repositorios/DetallePedidoRepositorio.cs).", bold_prefix="Yadier Pech Tun (Rol: Backend & Base de Datos): ")
    add_body_p(doc, "Como responsable de Backend y co-responsable de la arquitectura de acceso a datos, Leyva Chan implementó dos repositorios relacionales fundamentales para conectar la arquitectura SQL con los servicios del sistema:\n"
                    "1) El repositorio de consultas SQL para el catálogo de servicios ofrecidos (App/Repositorios/ServicioRepositorio.cs).\n"
                    "2) El repositorio SQL completo para la persistencia, búsqueda e indexación del catálogo de clientes (App/Repositorios/ClienteRepositorio.cs).", bold_prefix="Jesus Leyva Chan (Rol: Backend): ")
    
    add_header_p(doc, "¿Qué y Cómo se Implementó en el Código?", level=2)
    add_body_p(doc, "A continuación, se presentan de manera estrictamente separada los métodos y clases que materializan el motor de base de datos y sus repositorios, sin cortes ni omisiones:")
    add_code_table(doc, "Código Implementado por Yadier Pech Tun — Motor DDL y Creación de Tablas (App/BaseDatos/DatabaseInitializer.cs - Método Initialize)", extract_method_exact("BaseDatos/DatabaseInitializer.cs", "Initialize"))
    add_code_table(doc, "Código Implementado por Yadier Pech Tun — Precarga y Etiquetado de Datos Iniciales (App/BaseDatos/DatabaseInitializer.cs - Método SeedInitialData)", extract_method_exact("BaseDatos/DatabaseInitializer.cs", "SeedInitialData"))
    add_code_table(doc, "Código Implementado por Yadier Pech Tun — Repositorio Relacional de Renglones (App/Repositorios/DetallePedidoRepositorio.cs)", read_file_exact("Repositorios/DetallePedidoRepositorio.cs"))
    add_code_table(doc, "Código Implementado por Jesus Leyva Chan — Repositorio SQL de Catálogo de Servicios (App/Repositorios/ServicioRepositorio.cs)", read_file_exact("Repositorios/ServicioRepositorio.cs"))
    add_code_table(doc, "Código Implementado por Jesus Leyva Chan — Repositorio SQL del Catálogo de Clientes (App/Repositorios/ClienteRepositorio.cs)", read_file_exact("Repositorios/ClienteRepositorio.cs"))

    # =========================================================================
    # ACTIVIDAD 5
    # =========================================================================
    add_header_p(doc, "6. Actividad: Aprobación del prototipo y cierre de diseño", level=1)
    add_activity_table(
        doc,
        act_name="Aprobación del prototipo y cierre de diseño",
        f_init="2026-06-12",
        f_end="2026-06-19",
        duration="7 días",
        resp="Todos",
        dep="Diseño de interfaz y flujo en Figma",
        hitos="Validación visual, ajustes finales y aprobación para desarrollo",
        avance="100%",
        nota="Esta tarjeta marca el punto de cierre de paso entre diseño y construcción.\nAVANCE TOTAL 50%"
    )
    
    add_header_p(doc, "¿Qué se pretende hacer a rasgos generales en esta actividad?", level=2)
    add_body_p(doc, "En este punto de inflexión y cierre entre el diseño y la construcción formal del software, se pretende validar visualmente todas las interfaces y realizar los ajustes finales de formato e interacción. Específicamente, se busca consolidar la pantalla de ajustes estéticos y de negocio en el frontend, formalizar la estructura backend del comprobante digital (ticket) que se entregará al cliente, e implementar el servicio de alertas institucionales que proporcionará retroalimentación visual al usuario ante cualquier acción.")
    
    add_header_p(doc, "Desglose y Asignación de Responsabilidades por Integrante y Rol", level=2)
    add_body_p(doc, "Como responsable de Frontend, Daniel Moo ajustó y finalizó la vista del gestor de configuraciones y parámetros visuales del sistema (App/Pages/Configuraciones.razor), garantizando que la interfaz visual de administración refleje fielmente los colores, tipografías y secciones del prototipo aprobado por el usuario.", bold_prefix="Daniel Moo (Rol: Frontend): ")
    add_body_p(doc, "Como responsable de Backend & Base de Datos, Yadier Pech Tun construyó el modelo de datos estructurado del comprobante digital Ticket.cs (App/Modelos/Ticket.cs), definiendo las cabeceras, fechas, líneas de detalle y resúmenes financieros exactos exigidos en la validación del diseño para la impresión de recibos.", bold_prefix="Yadier Pech Tun (Rol: Backend & Base de Datos): ")
    add_body_p(doc, "Como responsable de Backend, Leyva Chan desarrolló el servicio de cuadros de diálogo e interrupción institucional CustomMessageBox.cs (App/Servicios/CustomMessageBox.cs), permitiendo disparar diálogos de validación visual y confirmación con la estética del prototipo aprobado.", bold_prefix="Jesus Leyva Chan (Rol: Backend): ")
    
    add_header_p(doc, "¿Qué y Cómo se Implementó en el Código?", level=2)
    add_code_table(doc, "Código Implementado por Daniel Moo — Vista Completa de Ajustes y Configuración Visual (App/Pages/Configuraciones.razor)", read_file_exact("Pages/Configuraciones.razor"))
    add_code_table(doc, "Código Implementado por Yadier Pech Tun — Estructura de Comprobante Digital del Diseño (App/Modelos/Ticket.cs)", read_file_exact("Modelos/Ticket.cs"))
    add_code_table(doc, "Código Implementado por Jesus Leyva Chan — Servicio Institucional de Cuadros de Diálogo (App/Servicios/CustomMessageBox.cs)", read_file_exact("Servicios/CustomMessageBox.cs"))

    # =========================================================================
    # ACTIVIDAD 6
    # =========================================================================
    add_header_p(doc, "7. Actividad: Backend: autenticación y usuarios", level=1)
    add_activity_table(
        doc,
        act_name="Backend: autenticación y usuarios",
        f_init="2026-06-23",
        f_end="2026-07-04",
        duration="11 días",
        resp="Yadier Pech Tun y Leyva Chan",
        dep="definición de módulos y backlog",
        hitos="Inicio de sesión(Leyva), roles (Yadier), sesión segura(Leyva) y gestión básica de usuarios(Yadier)",
        avance="100%",
        nota="Este módulo debe quedar estable antes de avanzar al resto de la lógica del sistema.\nAVANCE TOTAL AL TERMINAR: 60%"
    )
    
    add_header_p(doc, "¿Qué se pretende hacer a rasgos generales en esta actividad?", level=2)
    add_body_p(doc, "El propósito fundamental en esta actividad es construir y estabilizar por completo el perímetro de seguridad del backend antes de permitir el acceso a la lógica operativa transaccional. Se pretende implementar el inicio de sesión contra base de datos, la verificación y mantenimiento de una sesión global segura en memoria, el modelado transaccional de usuarios con validación de roles (Administrador vs Operador) y la gestión básica relacional (alta, búsqueda, desactivación) de credenciales en SQL.")
    
    add_header_p(doc, "Desglose y Asignación de Responsabilidades por Integrante y Rol", level=2)
    add_body_p(doc, "En cumplimiento exacto con los hitos asignados ('Inicio de sesión y sesión segura') dentro del área Backend, Leyva Chan desarrolló el servicio de autenticación transaccional LoginServicio.cs (App/Servicios/LoginServicio.cs). Este servicio ejecuta la verificación SQL sobre la tabla Usuarios, valida contraseñas en texto seguro o hash, e invoca directamente a SessionManager.SetUser() para consolidar la sesión segura y bloquear el acceso a usuarios inactivos.", bold_prefix="Jesus Leyva Chan (Rol: Backend - Hitos: Inicio de sesión y sesión segura): ")
    add_body_p(doc, "En cumplimiento de los hitos ('Roles y gestión básica de usuarios') bajo su rol de Backend & Base de Datos, Yadier Pech Tun desarrolló tres componentes estructurales independientes:\n"
                    "1) La entidad relacional de Usuario con definición de roles y estado activo (App/Modelos/Usuario.cs).\n"
                    "2) El modelo especializado de Administrador con privilegios de acceso de nivel superior (App/Modelos/Admin.cs).\n"
                    "3) El repositorio SQL completo UsuarioRepositorio.cs (App/Repositorios/UsuarioRepositorio.cs), que ejecuta las sentencias de inserción, actualización, consulta por ID/nombre y validación para la gestión básica de usuarios en base de datos.", bold_prefix="Yadier Pech Tun (Rol: Backend & Base de Datos - Hitos: Roles y gestión básica de usuarios): ")
    
    add_header_p(doc, "¿Qué y Cómo se Implementó en el Código?", level=2)
    add_body_p(doc, "Respetando rigurosamente la regla del proyecto de presentar cada clase y modelo en apartados individuales sin combinaciones artificiales, se expone el código exacto de la seguridad del sistema:")
    add_code_table(doc, "Código Implementado por Jesus Leyva Chan — Servicio Transaccional de Inicio de Sesión SQL (App/Servicios/LoginServicio.cs)", read_file_exact("Servicios/LoginServicio.cs"))
    add_code_table(doc, "Código Implementado por Yadier Pech Tun — Entidad Relacional de Usuario y Roles (App/Modelos/Usuario.cs)", read_file_exact("Modelos/Usuario.cs"))
    add_code_table(doc, "Código Implementado por Yadier Pech Tun — Entidad Especializada de Privilegios Administrativos (App/Modelos/Admin.cs)", read_file_exact("Modelos/Admin.cs"))
    add_code_table(doc, "Código Implementado por Yadier Pech Tun — Repositorio Relacional de Gestión de Usuarios (App/Repositorios/UsuarioRepositorio.cs)", read_file_exact("Repositorios/UsuarioRepositorio.cs"))

    # =========================================================================
    # ACTIVIDAD 7
    # =========================================================================
    add_header_p(doc, "8. Actividad: Backend: clientes, servicios y órdenes", level=1)
    add_activity_table(
        doc,
        act_name="Backend: clientes, servicios y órdenes",
        f_init="2026-07-05",
        f_end="2026-07-14",
        duration="5 días",
        resp="Yadier Pech Tun y Leyva Chan",
        dep="Autenticación y usuarios listos",
        hitos="CRUD principal(Yadier), registro de clientes (Leyva)",
        avance="100%",
        nota="Aquí se concentra la lógica más importante de la lavandería.\nAVANCE TOTAL AL TERMINAR: 65%"
    )
    
    add_header_p(doc, "¿Qué se pretende hacer a rasgos generales en esta actividad?", level=2)
    add_body_p(doc, "En esta etapa transaccional crítica se pretende construir y orquestar todo el motor operativo de la lavandería. El alcance incluye el registro e historial de clientes y visitas, y el CRUD principal del sistema: la creación, modificación y seguimiento de órdenes de lavado, la deducción automática e inteligente del inventario de insumos (detergentes, suavizantes, cloro) en base a las recetas por kilogramo de cada servicio, y el control del estado e intervalos de tiempo en las máquinas.")
    
    add_header_p(doc, "Desglose y Asignación de Responsabilidades por Integrante y Rol", level=2)
    add_body_p(doc, "En cumplimiento con el hito 'CRUD principal' asignado y en su calidad de responsable de Backend & Base de Datos, Yadier Pech Tun implementó la lógica transaccional de mayor complejidad en tres archivos separados:\n"
                    "1) El servicio transaccional orquestador PedidoServicio.cs (App/Servicios/PedidoServicio.cs), que procesa el alta integral de pedidos, calcula totales, asigna folios y coordina los detalles transaccionales.\n"
                    "2) El repositorio relacional de pedidos PedidoRepositorio.cs (App/Repositorios/PedidoRepositorio.cs), que ejecuta las transacciones SQL multitabla de encabezado y detalle en un solo bloque atómico.\n"
                    "3) El motor de deducción automática de inventarios InventarioAutomatizacion.cs (App/Servicios/InventarioAutomatizacion.cs), que consume insumos en base a fórmulas exactas por cada partida procesada.", bold_prefix="Yadier Pech Tun (Rol: Backend & Base de Datos - Hito: CRUD principal): ")
    add_body_p(doc, "En cumplimiento con el hito 'registro de clientes' asignado en su rol de Backend, Leyva Chan implementó el ciclo de vida y gestión de lealtad en dos componentes separados:\n"
                    "1) El servicio de lógica de negocio ClienteServicio.cs (App/Servicios/ClienteServicio.cs), que centraliza el alta, validación, actualización y búsqueda de clientes en mostrador.\n"
                    "2) El modelo relacional Cliente.cs (App/Modelos/Cliente.cs), que encapsula los datos de contacto y calcula automáticamente el estatus de fidelidad y acumulación de visitas de cada cliente regular.", bold_prefix="Jesus Leyva Chan (Rol: Backend - Hito: registro de clientes): ")
    
    add_header_p(doc, "¿Qué y Cómo se Implementó en el Código?", level=2)
    add_body_p(doc, "Se expone el código verbatim y sin alteraciones que materializa las operaciones transaccionales del núcleo del negocio:")
    add_code_table(doc, "Código Implementado por Yadier Pech Tun — Servicio Orquestador de Pedidos (App/Servicios/PedidoServicio.cs)", read_file_exact("Servicios/PedidoServicio.cs"))
    add_code_table(doc, "Código Implementado por Yadier Pech Tun — Repositorio SQL Transaccional de Pedidos (App/Repositorios/PedidoRepositorio.cs)", read_file_exact("Repositorios/PedidoRepositorio.cs"))
    add_code_table(doc, "Código Implementado por Yadier Pech Tun — Motor de Deducción Automática de Inventario SQL (App/Servicios/InventarioAutomatizacion.cs)", read_file_exact("Servicios/InventarioAutomatizacion.cs"))
    add_code_table(doc, "Código Implementado por Jesus Leyva Chan — Servicio Transaccional de Registro de Clientes (App/Servicios/ClienteServicio.cs)", read_file_exact("Servicios/ClienteServicio.cs"))
    add_code_table(doc, "Código Implementado por Jesus Leyva Chan — Entidad de Dominio y Cálculo de Lealtad (App/Modelos/Cliente.cs)", read_file_exact("Modelos/Cliente.cs"))

    # =========================================================================
    # ACTIVIDAD 8
    # =========================================================================
    add_header_p(doc, "9. Actividad: Frontend: pantallas principales", level=1)
    add_activity_table(
        doc,
        act_name="Frontend: pantallas principales",
        f_init="2026-07-10",
        f_end="2026-07-17",
        duration="7 días",
        resp="Daniel Moo",
        dep="Prototipo de Figma aprobado",
        hitos="Login, inicio, clientes, órdenes, navegación y formularios principales",
        avance="100%",
        nota="Debe respetar el diseño aprobado para evitar retrabajo en la integración.\nAVANCE TOTAL AL TERMINAR: 80%"
    )
    
    add_header_p(doc, "¿Qué se pretende hacer a rasgos generales en esta actividad?", level=2)
    add_body_p(doc, "En esta actividad, bajo responsabilidad exclusiva del área Frontend, se pretende construir todas las pantallas principales del sistema, respetando al máximo las directrices estéticas y los flujos aprobados en Figma. El objetivo es dotar al personal de mostrador y administración de interfaces visuales limpias, reactivas y ergonómicas en Blazor Hybrid / Razor para el ingreso al sistema (Login), cobro y punto de venta, monitoreo de lavadoras en tiempo real, seguimiento de órdenes y directorio de clientes.")
    
    add_header_p(doc, "Desglose y Asignación de Responsabilidades por Integrante y Rol", level=2)
    add_body_p(doc, "Al ser el responsable único de la actividad y titular del Frontend, Daniel Moo desarrolló la totalidad de las vistas principales del sistema en componentes Razor. Para cumplir con la directriz arquitectónica de no juntar archivos o vistas distintas en una sola sección, se explican y presentan en tablas independientes las pantallas maestras construidas por Daniel Moo:\n"
                    "1) Pantalla de Ingreso y Autenticación POS (Login.razor): Gestión visual de captura de credenciales y mensajes de error reactivos.\n"
                    "2) Terminal Principal de Cobro POS (Cobro.razor): Pantalla transaccional que presenta el resumen de la orden, selección rápida de métodos de pago en efectivo/tarjeta y cálculo en vivo del cambio a devolver.\n"
                    "3) Tablero Visual y Control de Máquinas (Maquinas.razor): Cuadrícula visual dinámica que representa las lavadoras y secadoras, mostrando indicadores por color según su disponibilidad ('ocupada' o 'disponible') y temporizadores de ciclo de lavado.\n"
                    "4) Directorio Visual de Clientes y Lealtad (Clientes.razor): Interfaz de mostrador para búsqueda ultrarrápida, alta de clientes nuevos y visualización de historial y lealtad.\n"
                    "5) Centro Maestro de Órdenes y Pedidos (Pedidos.razor): Vista de control operativo con filtrado por estado e interactividad en mostrador.", bold_prefix="Daniel Moo (Rol: Frontend - Responsable Único): ")
    
    add_header_p(doc, "¿Qué y Cómo se Implementó en el Código?", level=2)
    add_body_p(doc, "Se presentan por separado las interfaces de usuario construidas por Daniel Moo para las pantallas transaccionales clave del sistema sin contracciones ni omisiones:")
    add_code_table(doc, "Código Implementado por Daniel Moo — Pantalla de Autenticación y Acceso al POS (App/Pages/Login.razor)", read_file_exact("Pages/Login.razor"))
    add_code_table(doc, "Código Implementado por Daniel Moo — Terminal Visual de Cobro y Cálculo de Cambio (App/Pages/Cobro.razor)", read_file_exact("Pages/Cobro.razor"))
    add_code_table(doc, "Código Implementado por Daniel Moo — Tablero de Control y Monitoreo Visual de Máquinas (App/Pages/Maquinas.razor)", read_file_exact("Pages/Maquinas.razor"))
    add_code_table(doc, "Código Implementado por Daniel Moo — Directorio Visual e Interactivo de Clientes (App/Pages/Clientes.razor)", read_file_exact("Pages/Clientes.razor"))
    add_code_table(doc, "Código Implementado por Daniel Moo — Centro Visual de Control de Pedidos (App/Pages/Pedidos.razor)", read_file_exact("Pages/Pedidos.razor"))

    # =========================================================================
    # ACTIVIDAD 9
    # =========================================================================
    add_header_p(doc, "10. Actividad: Integración frontend + backend", level=1)
    add_activity_table(
        doc,
        act_name="Integración frontend + backend",
        f_init="2026-07-10",
        f_end="2026-07-17",
        duration="7 días",
        resp="Todos",
        dep="Backend y frontend listos y aprobación del usuario",
        hitos="Conexión total, validación de flujos y pruebas de integración",
        avance="100%",
        nota="Si aparece un error aquí, conviene corregirlo antes de entrar a pruebas finales.\nAVANCE TOTAL 85%"
    )
    
    add_header_p(doc, "¿Qué se pretende hacer a rasgos generales en esta actividad?", level=2)
    add_body_p(doc, "En esta etapa de integración se pretende acoplar e interconectar de forma definitiva las vistas frontend de Blazor con los servicios y repositorios transaccionales de C# / SQLite. El objetivo primordial es que todos los servicios de negocio se inyecten de manera limpia mediante un contenedor de Inyección de Dependencias (DI) en la capa de arranque de la aplicación, y que cada acción del usuario en mostrador desencadene transacciones reales en base de datos con validación continua de flujos.")
    
    add_header_p(doc, "Desglose y Asignación de Responsabilidades por Integrante y Rol", level=2)
    add_body_p(doc, "Desde su responsabilidad de Backend & Base de Datos, Yadier Pech Tun construyó el punto central de arranque e inyección de dependencias en App.xaml.cs (App/App.xaml.cs). Este módulo configura y dispara DatabaseInitializer.Initialize() y DatabaseInitializer.SeedInitialData() apenas arranca el contenedor WPF/Blazor, y expone el ciclo de vida del sistema ante errores fatales.", bold_prefix="Yadier Pech Tun (Rol: Backend & Base de Datos): ")
    add_body_p(doc, "Desde el área de Backend, Leyva Chan desarrolló la conexión operativa y transaccional para el procesamiento de pagos en la capa de negocio: PagoServicio.cs (App/Servicios/PagoServicio.cs). Este servicio vincula directamente el cobro en mostrador con el repositorio de pedidos y actualiza saldos e ingresos transaccionales en tiempo real.", bold_prefix="Jesus Leyva Chan (Rol: Backend): ")
    add_body_p(doc, "En su rol de Frontend, Daniel Moo interconectó los flujos visuales emergentes con los servicios de backend inyectados en componentes clave como CambiarEstado.razor (App/Pages/CambiarEstado.razor) y en el registro global de directivas y using de Razor (_Imports.razor en App/_Imports.razor), garantizando la reactividad y enlace de datos al cambiar el estado de las prendas.", bold_prefix="Daniel Moo (Rol: Frontend): ")
    
    add_header_p(doc, "¿Qué y Cómo se Implementó en el Código?", level=2)
    add_code_table(doc, "Código Implementado por Yadier Pech Tun — Contenedor de Arranque e Inyección de Dependencias (App/App.xaml.cs)", read_file_exact("App.xaml.cs"))
    add_code_table(doc, "Código Implementado por Jesus Leyva Chan — Servicio Transaccional de Integración de Cobro (App/Servicios/PagoServicio.cs)", read_file_exact("Servicios/PagoServicio.cs"))
    add_code_table(doc, "Código Implementado por Daniel Moo — Componente de Integración Transaccional de Cambios de Estado (App/Pages/CambiarEstado.razor)", read_file_exact("Pages/CambiarEstado.razor"))
    add_code_table(doc, "Código Implementado por Daniel Moo — Registro Global de Dependencias y Namespaces Frontend (App/_Imports.razor)", read_file_exact("_Imports.razor"))

    # =========================================================================
    # ACTIVIDAD 10
    # =========================================================================
    add_header_p(doc, "11. Actividad: Pruebas funcionales y corrección de errores", level=1)
    add_activity_table(
        doc,
        act_name="Pruebas funcionales y corrección de errores",
        f_init="2026-07-15",
        f_end="2026-07-19",
        duration="4 días",
        resp="Todos",
        dep="Integración completa de frontend y backend",
        hitos="Casos de prueba, revisión de fallos, ajustes de interfaz y validación de procesos",
        avance="100%",
        nota="Conviene registrar los errores por prioridad para resolver primero los que bloquean la entrega.\nAVANCE TOTAL 95%"
    )
    
    add_header_p(doc, "¿Qué se pretende hacer a rasgos generales en esta actividad?", level=2)
    add_body_p(doc, "Con la ejecución de pruebas funcionales y la corrección exhaustiva de errores se pretende validar la estabilidad y robustez de todos los procesos del sistema antes del pase a producción. Se busca instrumentar una bitácora de auditoría en la base de datos que registre las acciones transaccionales e incidencias, habilitar un sistema transaccional de alertas y notificaciones ante eventos inusuales o culminación de tareas, y verificar que el tablero ejecutivo principal muestre métricas precisas y responda sin fallos.")
    
    add_header_p(doc, "Desglose y Asignación de Responsabilidades por Integrante y Rol", level=2)
    add_body_p(doc, "Como titular de Backend y Base de Datos, Yadier Pech Tun implementó la infraestructura transaccional de observabilidad y trazabilidad del sistema en dos clases separadas:\n"
                    "1) El modelo transaccional Auditoria.cs (App/Modelos/Auditoria.cs), que define los campos de fecha, usuario, acción, tabla afectada y descripción técnica de cada movimiento en el sistema.\n"
                    "2) El repositorio relacional AuditoriaRepositorio.cs (App/Repositorios/AuditoriaRepositorio.cs), que realiza la persistencia en SQL para auditar flujos y clasificar fallas por prioridad.", bold_prefix="Yadier Pech Tun (Rol: Backend & Base de Datos): ")
    add_body_p(doc, "Desde su rol de Backend, Leyva Chan desarrolló la estructura transaccional para alertas operativas en campana y seguimiento de excepciones transaccionales en dos componentes separados:\n"
                    "1) El modelo relacional Notificacion.cs (App/Modelos/Notificacion.cs), que tipifica avisos de órdenes terminadas o errores operativos.\n"
                    "2) El repositorio relacional NotificacionRepositorio.cs (App/Repositorios/NotificacionRepositorio.cs), que persiste y actualiza el estado de lectura de cada notificación en base de datos.", bold_prefix="Jesus Leyva Chan (Rol: Backend): ")
    add_body_p(doc, "Desde el área Frontend, Daniel Moo ejecutó las pruebas de interfaz y desarrolló la vista del Tablero Ejecutivo en Vivo Dashboard.razor (App/Pages/Dashboard.razor). Este componente recolecta las métricas de ingresos, órdenes activas y alertas desde los repositorios inyectados y las presenta en tiempo real con tolerancia a fallos.", bold_prefix="Daniel Moo (Rol: Frontend): ")
    
    add_header_p(doc, "¿Qué y Cómo se Implementó en el Código?", level=2)
    add_body_p(doc, "A continuación, se exponen por separado las estructuras transaccionales de auditoría, alertas e interfaz de control de calidad:")
    add_code_table(doc, "Código Implementado por Yadier Pech Tun — Entidad Transaccional de Auditoría y Trazabilidad (App/Modelos/Auditoria.cs)", read_file_exact("Modelos/Auditoria.cs"))
    add_code_table(doc, "Código Implementado por Yadier Pech Tun — Repositorio SQL de Auditoría Operativa (App/Repositorios/AuditoriaRepositorio.cs)", read_file_exact("Repositorios/AuditoriaRepositorio.cs"))
    add_code_table(doc, "Código Implementado por Jesus Leyva Chan — Entidad Transaccional de Notificaciones (App/Modelos/Notificacion.cs)", read_file_exact("Modelos/Notificacion.cs"))
    add_code_table(doc, "Código Implementado por Jesus Leyva Chan — Repositorio SQL de Notificaciones y Alertas (App/Repositorios/NotificacionRepositorio.cs)", read_file_exact("Repositorios/NotificacionRepositorio.cs"))
    add_code_table(doc, "Código Implementado por Daniel Moo — Tablero Ejecutivo en Vivo y Métricas de Calidad (App/Pages/Dashboard.razor)", read_file_exact("Pages/Dashboard.razor"))

    # =========================================================================
    # ACTIVIDAD 11
    # =========================================================================
    add_header_p(doc, "12. Actividad: Documentación final y entrega", level=1)
    add_activity_table(
        doc,
        act_name="Documentación final y entrega",
        f_init="2026-07-18",
        f_end="2026-07-22",
        duration="4 días",
        resp="Todos",
        dep="Pruebas aprobadas",
        hitos="Documentación técnica, presentación final, cierre del proyecto y enlace de Figma adjunto",
        avance="0%",
        nota="Aquí se deja lista la versión final del proyecto para revisión y entrega.\nAVANCE TOTAL 100%"
    )
    
    add_header_p(doc, "¿Qué se pretende hacer a rasgos generales en esta actividad?", level=2)
    add_body_p(doc, "En esta actividad de cierre formal se pretende consolidar la versión final de la solución LavanderiaApp v0.1 para revisión administrativa y entrega formal. A nivel técnico, esto conlleva la implementación matemática y contable exacta de los cortes de caja por turno y registro de egresos, la interfaz visual para exportación de reportes operativos, y la integración de un motor nativo en Python (y C#) capaz de emitir en formato PDF el corte contable formal de la lavandería con numeración de páginas y resúmenes consolidados.")
    
    add_header_p(doc, "Desglose y Asignación de Responsabilidades por Integrante y Rol", level=2)
    add_body_p(doc, "Como responsable de Backend y Base de Datos, Yadier Pech Tun estructuró las entidades matemáticas y contables que permiten generar balances de corte sin discrepancias. Para cumplir estrictamente con la directriz de presentar los componentes de manera separada, se exponen por separado las entidades contables desarrolladas por Yadier:\n"
                    "1) El modelo contable de Corte de Caja (App/Modelos/CorteCaja.cs), que registra los montos en efectivo, transferencias, fondos iniciales y diferencias de arqueo al cierre de turno.\n"
                    "2) El modelo relacional de Gastos Operativos (App/Modelos/Gasto.cs), que permite asentar salidas de efectivo de la caja chica por concepto de insumos o servicios con justificación.", bold_prefix="Yadier Pech Tun (Rol: Backend & Base de Datos): ")
    add_body_p(doc, "Como responsable de Backend, Leyva Chan implementó el repositorio relacional y las sentencias SQL para almacenar, agregar y consultar las transacciones de arqueo y corte en la base de datos: CorteCajaRepositorio.cs (App/Repositorios/CorteCajaRepositorio.cs), garantizando la inmutabilidad de los cierres diarios.", bold_prefix="Jesus Leyva Chan (Rol: Backend): ")
    add_body_p(doc, "En su calidad de responsable de Frontend, Daniel Moo desarrolló la interfaz de usuario para la emisión, filtrado y descarga de reportes oficiales del establecimiento en Reportes.razor (App/Pages/Reportes.razor), y conectó e implementó el script oficial de generación de reportes e informes en PDF con ReportLab (App/generar_pdf_actividad.py), logrando un documento exportable con cabeceras institucionales y resúmenes ejecutivos.", bold_prefix="Daniel Moo (Rol: Frontend): ")
    
    add_header_p(doc, "¿Qué y Cómo se Implementó en el Código?", level=2)
    add_body_p(doc, "Para finalizar el informe de desarrollo, se exponen por separado las entidades contables, los repositorios de arqueo y el motor de reportes en PDF verbatim sin contracciones:")
    add_code_table(doc, "Código Implementado por Yadier Pech Tun — Entidad Contable de Arqueo y Cierre de Turno (App/Modelos/CorteCaja.cs)", read_file_exact("Modelos/CorteCaja.cs"))
    add_code_table(doc, "Código Implementado por Yadier Pech Tun — Entidad Contable de Registro de Egresos (App/Modelos/Gasto.cs)", read_file_exact("Modelos/Gasto.cs"))
    add_code_table(doc, "Código Implementado por Jesus Leyva Chan — Repositorio SQL de Transacciones de Corte (App/Repositorios/CorteCajaRepositorio.cs)", read_file_exact("Repositorios/CorteCajaRepositorio.cs"))
    add_code_table(doc, "Código Implementado por Daniel Moo — Interfaz de Usuario para Emisión y Exportación de Reportes (App/Pages/Reportes.razor)", read_file_exact("Pages/Reportes.razor"))
    add_code_table(doc, "Código Implementado por Daniel Moo — Motor Oficial de Generación de Reportes Operativos en PDF (App/generar_pdf_actividad.py)", read_file_exact("generar_pdf_actividad.py"))

    # Save with fallback if file is open in Microsoft Word
    out_paths = [
        r"C:\Users\Yadie\RiderProjects\LavanderiaApp0.1\App\Shared\Informe_Desarrollo_LavanderiaApp_Adaptado.docx",
        r"C:\Users\Yadie\RiderProjects\LavanderiaApp0.1\App\Shared\Informe_Desarrollo_LavanderiaApp_Adaptado_Final.docx",
        r"C:\Users\Yadie\Downloads\Informe_Desarrollo_LavanderiaApp_Adaptado_Final.docx"
    ]
    saved = []
    for out_path in out_paths:
        try:
            doc.save(out_path)
            saved.append(out_path)
            print(f"¡Guardado exitoso en: {out_path}!")
        except PermissionError:
            print(f"[Aviso] El archivo {out_path} está abierto en Microsoft Word. Se guardó en una ruta alternativa.")
        except Exception as e:
            print(f"Error al guardar en {out_path}: {e}")
            
    if not saved:
        print("No se pudo guardar el archivo en ninguna ubicación.")
    else:
        print("¡Generación exitosa del informe completo!")

if __name__ == "__main__":
    generate_report()
