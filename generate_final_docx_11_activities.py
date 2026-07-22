import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table, color="D3D3D3", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), val)
        border.set(qn('w:sz'), sz)
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tblBorders.append(border)
    border = OxmlElement('w:insideV')
    border.set(qn('w:val'), 'none')
    tblBorders.append(border)
    tblPr.append(tblBorders)

def add_header_p(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Segoe UI'
    run.bold = True
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(30, 58, 138) # #1E3A8A
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(15, 23, 42) # #0F172A
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
    elif level == 3:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(2, 132, 199) # #0284C7
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
    elif level == 4:
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(16, 185, 129) # Verde esmeralda para Hito
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
    return p

def add_body_p(doc, text, bold_prefix=None, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = 'Segoe UI'
        r_pre.bold = True
        r_pre.font.size = Pt(10.5)
        r_pre.font.color.rgb = RGBColor(15, 23, 42)
    run = p.add_run(text)
    run.font.name = 'Segoe UI'
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(51, 65, 85)
    run.italic = italic
    return p

def add_activity_table(doc, act_name, f_init, f_end, duration, resp, dep, hitos, avance, nota):
    table = doc.add_table(rows=8, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table, color="CBD5E1", sz="6")
    
    col_widths = [Inches(2.2), Inches(4.3)]
    for row in table.rows:
        for idx, width in enumerate(col_widths):
            row.cells[idx].width = width
            
    headers = [
        "Identificación de la Actividad",
        "Fecha de Inicio / Fin",
        "Duración Estimada",
        "Responsable(s)",
        "Dependencia",
        "Hitos",
        "Porcentaje de Avance",
        "Nota / Base del Proyecto"
    ]
    values = [
        act_name,
        f"{f_init}  al  {f_end}",
        duration,
        resp,
        dep,
        hitos,
        avance,
        nota
    ]
    
    for i, (h, v) in enumerate(zip(headers, values)):
        row = table.rows[i]
        c0, c1 = row.cells[0], row.cells[1]
        set_cell_margins(c0, top=100, bottom=100, left=150, right=150)
        set_cell_margins(c1, top=100, bottom=100, left=150, right=150)
        
        if i == 0:
            set_cell_background(c0, "1E3A8A")
            set_cell_background(c1, "1E3A8A")
            p0 = c0.paragraphs[0]
            r0 = p0.add_run(h)
            r0.font.name = 'Segoe UI'
            r0.bold = True
            r0.font.size = Pt(10)
            r0.font.color.rgb = RGBColor(255, 255, 255)
            
            p1 = c1.paragraphs[0]
            r1 = p1.add_run(v)
            r1.font.name = 'Segoe UI'
            r1.bold = True
            r1.font.size = Pt(10)
            r1.font.color.rgb = RGBColor(255, 255, 255)
        else:
            bg = "F8FAFC" if i % 2 == 1 else "FFFFFF"
            set_cell_background(c0, bg)
            set_cell_background(c1, bg)
            
            p0 = c0.paragraphs[0]
            r0 = p0.add_run(h)
            r0.font.name = 'Segoe UI'
            r0.bold = True
            r0.font.size = Pt(9.5)
            r0.font.color.rgb = RGBColor(30, 41, 59)
            
            p1 = c1.paragraphs[0]
            r1 = p1.add_run(v)
            r1.font.name = 'Segoe UI'
            r1.font.size = Pt(9.5)
            if i == 6:
                r1.bold = True
                r1.font.color.rgb = RGBColor(2, 132, 199)
            else:
                r1.font.color.rgb = RGBColor(51, 65, 85)
    
    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_before = Pt(4)
    p_space.paragraph_format.space_after = Pt(4)

def read_file_exact(rel_path):
    base_dir = r"C:\Users\Yadie\RiderProjects\LavanderiaApp0.1\App"
    full_path = os.path.join(base_dir, rel_path)
    if not os.path.exists(full_path):
        return f"// ERROR: No se encontró el archivo {rel_path}"
    with open(full_path, "r", encoding="utf-8") as f:
        return f.read().rstrip()

def read_file_lines_exact(rel_path, start_line, end_line):
    base_dir = r"C:\Users\Yadie\RiderProjects\LavanderiaApp0.1\App"
    full_path = os.path.join(base_dir, rel_path)
    if not os.path.exists(full_path):
        return f"// ERROR: No se encontró el archivo {rel_path}"
    with open(full_path, "r", encoding="utf-8") as f:
        lines = f.readlines()
    s = max(0, start_line - 1)
    e = min(len(lines), end_line)
    return "".join(lines[s:e]).rstrip()

def extract_method_exact(rel_path, method_name, start_keyword="public"):
    base_dir = r"C:\Users\Yadie\RiderProjects\LavanderiaApp0.1\App"
    full_path = os.path.join(base_dir, rel_path)
    with open(full_path, "r", encoding="utf-8") as f:
        lines = f.readlines()
    
    start_idx = -1
    for i, l in enumerate(lines):
        if method_name in l and (start_keyword in l or "void" in l or "Task" in l or "List" in l or "bool" in l):
            start_idx = i
            break
            
    if start_idx == -1:
        return read_file_exact(rel_path)
        
    result_lines = []
    brace_count = 0
    found_first_brace = False
    
    for i in range(start_idx, len(lines)):
        l = lines[i]
        result_lines.append(l)
        if "{" in l:
            brace_count += l.count("{")
            found_first_brace = True
        if "}" in l:
            brace_count -= l.count("}")
        if found_first_brace and brace_count == 0:
            break
            
    return "".join(result_lines).rstrip()

def add_code_table_with_explanation(doc, title, code_text, hito_explanation):
    add_header_p(doc, title, level=3)
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table, color="94A3B8", sz="6")
    
    cell = table.rows[0].cells[0]
    cell.width = Inches(6.5)
    set_cell_background(cell, "0F172A") # dark slate theme
    set_cell_margins(cell, top=140, bottom=140, left=160, right=160)
    
    p = cell.paragraphs[0]
    p.paragraph_format.line_spacing = 1.05
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(241, 245, 249) # light text
    
    add_header_p(doc, "Explicación Detallada del Hito Hecho:", level=4)
    add_body_p(doc, hito_explanation, italic=False)

def generate_full_report():
    doc = docx.Document()
    
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        
    p_title = doc.add_paragraph()
    r_t = p_title.add_run("SISTEMA LAVANDERIAAPP v0.1\nInforme Técnico Exhaustivo de Desarrollo por Actividades, Roles Institucionales, Código Verbatim y Validación de Hitos")
    r_t.font.name = 'Segoe UI'
    r_t.font.size = Pt(20)
    r_t.bold = True
    r_t.font.color.rgb = RGBColor(30, 58, 138)
    p_title.paragraph_format.space_after = Pt(14)
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    add_header_p(doc, "1. Introducción y Asignación de Roles del Equipo", level=1)
    add_body_p(doc, "El presente documento constituye el informe final y la memoria técnica integral de desarrollo del software LavanderiaApp v0.1. En estricto cumplimiento con la metodología de gestión por actividades y tarjetas de control de proyecto, cada uno de los 11 apartados se desarrolla bajo un protocolo estructurado en cuatro fases obligatorias:\n"
                    "1. Identificación y Ficha de la Actividad: Resumen de fechas, responsables, dependencias, hitos y porcentaje de avance.\n"
                    "2. Explicación General de la Actividad: Detalle preciso de lo que se pretende hacer y el objetivo técnico de la fase.\n"
                    "3. Desglose de Responsabilidades y Explicación por Integrante y Rol: Análisis individual de la labor realizada por cada integrante asignado al apartado según su especialidad.\n"
                    "4. Código Referenciado Verbatim y Explicación del Hito Hecho: Inserción del código original 100% exacto (sin contracciones, sin recortes y separando partes del mismo archivo en bloques independientes) seguido de la explicación técnica particular de cada hito alcanzado.")
    
    add_body_p(doc, "• Yadier Pech Tun (Rol: Backend & Base de Datos): Responsable integral del diseño y arquitectura relacional SQLite, modelado DDL, repositorios transaccionales SQL, automatización matemática de inventarios, lógica operativa de pedidos y trazabilidad por auditoría.", bold_prefix="Rol 1: ")
    add_body_p(doc, "• Jesus Leyva Chan / Leyva Chan (Rol: Backend): Responsable integral de la seguridad transaccional (autenticación, control de sesiones, cifrado), servicios relacionales de clientes y lealtad, motor de cobros/pagos y sistema de alertas y notificaciones.", bold_prefix="Rol 2: ")
    add_body_p(doc, "• Daniel Moo (Rol: Frontend): Responsable exclusivo de la interfaz gráfica y experiencia de usuario, desarrollando las pantallas del Punto de Venta (POS), navegación en Blazor Hybrid / Razor, tableros de control de máquinas y vistas de reportes.", bold_prefix="Rol 3: ")

    # =========================================================================
    # ACTIVIDAD 1
    # =========================================================================
    add_header_p(doc, "2. Actividad: Levantamiento de requerimientos y definición del alcance", level=1)
    add_activity_table(
        doc,
        act_name="Levantamiento de requerimientos y definición del alcance",
        f_init="2026-05-07",
        f_end="2026-06-02",
        duration="29 días",
        resp="Yadier Pech Tun, Leyva Chan y Daniel Moo",
        dep="Ninguna",
        hitos="Lista de módulos, objetivos del sistema, alcance validado y prioridades del proyecto",
        avance="100%",
        nota="De aquí sale la base del resto del proyecto: módulos, reglas de negocio y entregables mínimos."
    )
    
    add_header_p(doc, "¿Qué se pretende hacer a rasgos generales en esta actividad?", level=2)
    add_body_p(doc, "En esta actividad inicial se pretende sentar los cimientos conceptuales, normativos y estructurales sobre los cuales operará toda la plataforma LavanderiaApp v0.1. El objetivo técnico central es transformar las necesidades reales del negocio de lavandería en especificaciones de software concretas: establecer los módulos primarios (Usuarios, Clientes, Servicios, Pedidos, Máquinas e Inventario), fijar las reglas de negocio inmutables (tasa de IVA del 16%, fondo de caja inicial y datos corporativos) y configurar el marco de sesión y presentación visual que regulará el acceso y la experiencia de los operadores.")
    
    add_header_p(doc, "Explicación de Cada Integrante con su Respectivo Rol", level=2)
    add_body_p(doc, "Como responsable del área de Backend & Base de Datos, Yadier Pech Tun implementó la centralización de los parámetros operativos y fiscales del negocio en la clase estática BusinessConfig.cs. De esta forma garantizó que las reglas monetarias base (tasa impositiva del 16%, fondo inicial del mostrador y datos de cabecera) no se encuentren hardcodeadas o dispersas, sino accesibles desde una única fuente de verdad transaccional que alimenta a los repositorios y servicios de base de datos.", bold_prefix="Yadier Pech Tun (Rol: Backend & Base de Datos): ")
    add_body_p(doc, "Como responsable del Backend de seguridad, Leyva Chan construyó el gestor global de sesión SessionManager.cs. Su labor en esta fase consistió en delimitar el alcance del sistema en memoria, estableciendo el contenedor que almacena la identidad del usuario activo y expone validaciones inmediatas de nivel de privilegio (Usuario vs Admin), impidiendo que transacciones operativas se ejecuten sin respaldo de una sesión válida.", bold_prefix="Jesus Leyva Chan (Rol: Backend): ")
    add_body_p(doc, "Como responsable de Frontend, Daniel Moo estructuró el esqueleto visual del contenedor POS en MainLayout.razor. Para cumplir estrictamente con la regla de presentar las partes de un mismo archivo por separado sin juntarlas, Daniel separó la estructura del menú lateral y cabecera visual (líneas 1 a 485) de la lógica transaccional y de notificaciones en código C# (líneas 486 a 1145), materializando las prioridades de navegación validadas en el alcance.", bold_prefix="Daniel Moo (Rol: Frontend): ")
    
    add_header_p(doc, "¿Qué y Cómo se Implementó en el Código? (Referencia Verbatim y Explicación por Hito)", level=2)
    
    add_code_table_with_explanation(
        doc,
        title="Código Implementado por Yadier Pech Tun — Reglas de Negocio y Configuración (App/Servicios/BusinessConfig.cs)",
        code_text=read_file_exact("Servicios/BusinessConfig.cs"),
        hito_explanation="Este hito materializa la 'Lista de módulos, objetivos del sistema y alcance validado'. Al definir los parámetros fiscales (IVA 16%), el fondo de caja obligatorio e información de contacto centralizada, el código de BusinessConfig.cs actúa como el cimiento operativo del backend. Garantiza que todos los cálculos financieros de pedidos y cortes de caja que se realizan en etapas posteriores utilicen exactamente estas reglas de negocio unificadas."
    )
    
    add_code_table_with_explanation(
        doc,
        title="Código Implementado por Jesus Leyva Chan — Gestión del Alcance de Sesión Segura (App/Servicios/SessionManager.cs)",
        code_text=read_file_exact("Servicios/SessionManager.cs"),
        hito_explanation="Este hito cumple con la definición del alcance de seguridad del usuario en memoria. La clase estática SessionManager resguarda la entidad Usuario que ha iniciado sesión (`UsuarioActual`) y provee los métodos de interrogación `IsAdmin()` e `IsLoggedIn()`. Esto asegura desde el backend que el alcance operativo esté delimitado: un usuario sin sesión o sin rol de administrador es bloqueado antes de invocar cualquier transacción en la base de datos."
    )
    
    add_code_table_with_explanation(
        doc,
        title="Código Implementado por Daniel Moo — Estructura y Navegación del POS (App/Shared/MainLayout.razor - Sección Interfaz Visual)",
        code_text=read_file_lines_exact("Shared/MainLayout.razor", 1, 485),
        hito_explanation="Esta primera parte del código de MainLayout.razor representa el hito de 'Prioridades del proyecto y navegación visual'. Define el menú lateral (Sidebar) con accesos condicionados por los permisos del usuario activo, la barra superior con indicador de estado y la estructura de contenedores que alojan el resto de las vistas de la aplicación Blazor."
    )
    
    add_code_table_with_explanation(
        doc,
        title="Código Implementado por Daniel Moo — Lógica del Contenedor y Alertas (App/Shared/MainLayout.razor - Sección Lógica C# @code)",
        code_text=read_file_lines_exact("Shared/MainLayout.razor", 486, 1145),
        hito_explanation="Esta segunda parte independiente de MainLayout.razor (el bloque `@code`) implementa la lógica de interacción del contenedor maestro. Gestiona el temporizador de notificaciones tipo Toast, el buscador global en vivo de pedidos/clientes y la interconexión con los modales rápidos de corte de caja y registro de gastos operativos desde el mostrador."
    )

    # =========================================================================
    # ACTIVIDAD 2
    # =========================================================================
    add_header_p(doc, "3. Actividad: Definición de backlog y módulos principales", level=1)
    add_activity_table(
        doc,
        act_name="Definición de backlog y módulos principales",
        f_init="2026-06-07",
        f_end="2026-06-11",
        duration="4 días",
        resp="Yadier Pech Tun",
        dep="Requerimientos definidos",
        hitos="Módulos de usuarios, clientes, servicios, pedidos y reportes detalles pedido, etc",
        avance="100%",
        nota="Apartado ya listo\nAVANCE TOTAL AL TERMINAR: 25%"
    )
    
    add_header_p(doc, "¿Qué se pretende hacer a rasgos generales en esta actividad?", level=2)
    add_body_p(doc, "En este apartado, bajo la responsabilidad única y exclusiva de Yadier Pech Tun (Backend & Base de Datos), se pretende estructurar conceptual y físicamente el modelo de dominio y el backlog transaccional del sistema. El objetivo es definir las clases de entidad en C# para los módulos operacionales críticos (Pedidos y Servicios) fijando sus propiedades, estados de flujo y reglas matemáticas de cálculo en memoria, y construir el módulo resolutor de rutas para conectar la capa transaccional con la base de datos SQLite.")
    
    add_header_p(doc, "Explicación de Cada Integrante con su Respectivo Rol", level=2)
    add_body_p(doc, "Al ser el responsable único del apartado y titular del área de Base de Datos y Backend, Yadier Pech Tun conceptualizó, programó y validó de forma íntegra las tres entidades de dominio y configuración que sustentan el backlog del proyecto:\n"
                    "1) El módulo de resolución de rutas y alertas Config.cs, que localiza dinámicamente el archivo físico SQLite lavanderia.db y expone el observador de notificaciones transaccionales ToastService.\n"
                    "2) La entidad central del módulo de Pedidos (Pedido.cs), que define la estructura completa de una orden: autogeneración de folios, relación con cliente, fechas de promesa, máquina de estados de las prendas ('En espera', 'Lavando', 'Secando', 'Listo', 'Entregado') y cálculo automático de saldos pendientes, anticipos y totales a pagar.\n"
                    "3) La entidad del catálogo comercial (Servicio.cs), donde se programaron las tarifas por kilogramo o pieza, los tiempos estimados de ciclo y los métodos matemáticos para calcular el importe por cantidad.", bold_prefix="Yadier Pech Tun (Rol: Backend & Base de Datos - Responsable Único): ")
    add_body_p(doc, "Para mantener una trazabilidad institucional clara, se hace constar que en este apartado técnico de definición conceptual de modelos y backlog transaccional, Leyva Chan (Backend) y Daniel Moo (Frontend) no cuentan con asignación directa en la tarjeta, dado que es una labor especializada de diseño de base de datos y entidades por parte de Yadier Pech Tun.", bold_prefix="Jesus Leyva Chan y Daniel Moo: ")
    
    add_header_p(doc, "¿Qué y Cómo se Implementó en el Código? (Referencia Verbatim y Explicación por Hito)", level=2)
    
    add_code_table_with_explanation(
        doc,
        title="Código Implementado por Yadier Pech Tun — Resolución de Rutas DB y Bus de Alertas (App/BaseDatos/Config.cs)",
        code_text=read_file_exact("BaseDatos/Config.cs"),
        hito_explanation="Este hito asegura la conectividad y ubicación física de la base de datos para todos los módulos definidos. La clase Config resuelve de forma autónoma el directorio de execution y la ruta del archivo `lavanderia.db`, al tiempo que expone el evento estático `ToastService` que permite a los repositorios notificar el éxito o fracaso de las transacciones al frontend."
    )
    
    add_code_table_with_explanation(
        doc,
        title="Código Implementado por Yadier Pech Tun — Entidad Central del Módulo de Pedidos (App/Modelos/Pedido.cs)",
        code_text=read_file_exact("Modelos/Pedido.cs"),
        hito_explanation="Este archivo representa el hito más importante del backlog operacional: la entidad Pedido. Implementa el ciclo de vida transaccional completo del negocio, incluyendo enumeradores para estados operativos, cálculo en tiempo real del saldo pendiente (`Total - Anticipo`) e identificación única por folio y código de cliente."
    )
    
    add_code_table_with_explanation(
        doc,
        title="Código Implementado por Yadier Pech Tun — Entidad del Catálogo de Servicios (App/Modelos/Servicio.cs)",
        code_text=read_file_exact("Modelos/Servicio.cs"),
        hito_explanation="Este hito completa el backlog transaccional al definir el catálogo comercial de la lavandería. La clase Servicio encapsula el precio unitario (por kilogramo, pieza o carga), el tiempo de entrega en horas y el método formal de cálculo de subtotal que será invocado cada vez que se agregue una prenda a la orden."
    )

    # =========================================================================
    # ACTIVIDAD 3
    # =========================================================================
    add_header_p(doc, "4. Actividad: Diseño de interfaz y flujo en Figma", level=1)
    add_activity_table(
        doc,
        act_name="Diseño de interfaz y flujo en Figma",
        f_init="2026-06-05",
        f_end="2026-06-07",
        duration="2 días",
        resp="Yadier Pech Tun, Jesus Leyva Chan, Daniel Moo",
        dep="Requerimientos definidos",
        hitos="Wireframes, flujo de pantallas, prototipo navegable y estructura visual",
        avance="100%",
        nota="AVANCE TOTAL 35%"
    )
    
    add_header_p(doc, "¿Qué se pretende hacer a rasgos generales en esta actividad?", level=2)
    add_body_p(doc, "El objetivo primordial en este apartado es trasladar los wireframes, flujos de pantalla y prototipos navegables concebidos en Figma hacia estructuras concretas dentro del proyecto en C# y Blazor. Se pretende que el mostrador cuente con componentes visuales modales ágiles para la selección del método de cobro (Efectivo vs Tarjeta) y que el backend disponga en paralelo de los modelos exactos para representar cada renglón o ítem seleccionado en mostrador antes y durante su registro transaccional en la base de datos.")
    
    add_header_p(doc, "Explicación de Cada Integrante con su Respectivo Rol", level=2)
    add_body_p(doc, "Como responsable de Frontend, Daniel Moo programó el componente modal interactivo TipoPago.razor, convirtiendo el wireframe de selección de pago de Figma en un diálogo Razor reactivo. Este módulo permite al operador elegir entre pago en Efectivo o Tarjeta, redimensionando la interfaz para mostrar el cálculo en vivo de cambio a devolver o la confirmación de terminal electrónica.", bold_prefix="Daniel Moo (Rol: Frontend): ")
    add_body_p(doc, "Como responsable de Backend & Base de Datos, Yadier Pech Tun construyó la entidad relacional de renglón Detalle_Pedido.cs, la cual actúa como el soporte estructurado en base de datos para cada línea visualizada en las tablas del prototipo, multiplicando peso por tarifa y consolidando el importe por partida.", bold_prefix="Yadier Pech Tun (Rol: Backend & Base de Datos): ")
    add_body_p(doc, "Como responsable de Backend, Leyva Chan desarrolló el modelo transaccional en memoria CarritoItem.cs. Su aportación resolvió el flujo visual en mostrador: mientras el usuario navega por las pantallas y hace clic en los servicios del catálogo, CarritoItem mantiene el desglose temporal en la memoria RAM antes de persistirlo de forma definitiva en SQL.", bold_prefix="Jesus Leyva Chan (Rol: Backend): ")
    
    add_header_p(doc, "¿Qué y Cómo se Implementó en el Código? (Referencia Verbatim y Explicación por Hito)", level=2)
    
    add_code_table_with_explanation(
        doc,
        title="Código Implementado por Daniel Moo — Modal Interactivo de Selección de Pago (App/Pages/TipoPago.razor)",
        code_text=read_file_exact("Pages/TipoPago.razor"),
        hito_explanation="Este componente Razor cumple con el hito de 'Prototipo navegable y flujo de pantallas de cobro'. Ofrece una interfaz modal con tarjetas interactivas y retroalimentación táctil/visual inmediata para la selección de método de pago, redirigiendo automáticamente la pantalla al terminal de cobro correspondiente con los parámetros exactos de la orden en curso."
    )
    
    add_code_table_with_explanation(
        doc,
        title="Código Implementado por Yadier Pech Tun — Soporte Relacional de Renglones (App/Modelos/Detalle_Pedido.cs)",
        code_text=read_file_exact("Modelos/Detalle_Pedido.cs"),
        hito_explanation="Este hito representa la estructura visual y de datos en la base de datos para el desglose del pedido. Cada instancia de Detalle_Pedido almacena la referencia al servicio, la cantidad de kilogramos o piezas, el precio unitario con el que se vendió y calcula el subtotal exacto de esa línea."
    )
    
    add_code_table_with_explanation(
        doc,
        title="Código Implementado por Jesus Leyva Chan — Modelo del Carrito de Mostrador en Memoria (App/Modelos/CarritoItem.cs)",
        code_text=read_file_exact("Modelos/CarritoItem.cs"),
        hito_explanation="Este hito materializa el flujo dinámico del mostrador en la capa intermedia. La clase CarritoItem encapsula de manera ligera cada servicio agregado desde la pantalla del POS, calculando al instante el subtotal en memoria para actualizar el resumen visual antes de ejecutar la transacción SQL."
    )

    # =========================================================================
    # ACTIVIDAD 4
    # =========================================================================
    add_header_p(doc, "5. Actividad: Diseño de base de datos y arquitectura", level=1)
    add_activity_table(
        doc,
        act_name="Diseño de base de datos y arquitectura",
        f_init="2026-06-07",
        f_end="2026-06-11",
        duration="4 días",
        resp="Yadier Pech Tun y Jesus Leyva Chan",
        dep="definición de backlogs y módulos",
        hitos="Modelo entidad-relación, estructura de tablas, SQLite con reglas base y etiquetas",
        avance="25%",
        nota="De este diseño dependerán autenticación, clientes, servicios, órdenes y reportes.\nAVANCE TOTAL 40%"
    )
    
    add_header_p(doc, "¿Qué se pretende hacer a rasgos generales en esta actividad?", level=2)
    add_body_p(doc, "En este apartado técnico crítico se pretende construir el motor de persistencia relacional SQLite y diseñar la arquitectura de acceso a datos de la aplicación. El objetivo general es ejecutar sentencias DDL estrictas que creen las tablas del modelo entidad-relación (con llaves primarias, foráneas autovalidadas y reglas predeterminadas), inyectar los datos y etiquetas del catálogo inicial (Seeding) y programar los repositorios SQL que permitirán realizar operaciones transaccionales de lectura y escritura sobre renglones de pedido, servicios y clientes.")
    
    add_header_p(doc, "Explicación de Cada Integrante con su Respectivo Rol", level=2)
    add_body_p(doc, "Como líder indiscutible del área de Base de Datos y Backend, Yadier Pech Tun construyó el motor transaccional DatabaseInitializer.cs. Para respetar de forma rigurosa la norma del proyecto de no juntar partes distintas del mismo archivo en una sola tabla, Yadier desglosó su implementación en dos bloques independientes:\n"
                    "1) El método Initialize(), que ejecuta la activación de PRAGMA foreign_keys = ON y las sentencias DDL CREATE TABLE para todas las tablas relacionales maestras.\n"
                    "2) El método SeedInitialData(), que inyecta de forma atómica el catálogo inicial y las etiquetas obligatorias (como la cuenta de superadministrador con hash y la lista de servicios estándar).\n"
                    "Adicionalmente, Yadier implementó el repositorio relacional transaccional DetallePedidoRepositorio.cs para persistir los ítems de las órdenes.", bold_prefix="Yadier Pech Tun (Rol: Backend & Base de Datos): ")
    add_body_p(doc, "Como responsable de Backend y co-responsable de la arquitectura de datos, Leyva Chan desarrolló dos repositorios relacionales fundamentales para interconectar el modelo SQL con los servicios en C#:\n"
                    "1) ServicioRepositorio.cs, encargado de realizar las consultas SQL de lectura, búsqueda y edición sobre el catálogo de servicios de lavandería.\n"
                    "2) ClienteRepositorio.cs, responsable de indexar, consultar, registrar y sumar visitas acumuladas en la tabla de clientes de SQLite.", bold_prefix="Jesus Leyva Chan (Rol: Backend): ")
    add_body_p(doc, "Dado que esta tarjeta se enfoca estrictamente en la arquitectura del motor SQL relacional, sentencias DDL y repositorios transaccionales, Daniel Moo (Frontend) no cuenta con asignación directa en esta fase, concentrando sus esfuerzos en las capas de presentación visual que consumirán este motor.", bold_prefix="Daniel Moo (Rol: Frontend): ")
    
    add_header_p(doc, "¿Qué y Cómo se Implementó en el Código? (Referencia Verbatim y Explicación por Hito)", level=2)
    
    add_code_table_with_explanation(
        doc,
        title="Código Implementado por Yadier Pech Tun — Motor DDL y Creación de Tablas (App/BaseDatos/DatabaseInitializer.cs - Método Initialize)",
        code_text=extract_method_exact("BaseDatos/DatabaseInitializer.cs", "Initialize"),
        hito_explanation="Este método cumple con el hito de 'Modelo entidad-relación y estructura de tablas'. Ejecuta la inicialización física en el motor SQLite, asegurando que las llaves foráneas estén activas para evitar datos huérfanos y creando con tipos estrictos las tablas de Usuarios, Clientes, Servicios, Pedidos, DetallesPedido, Auditoría, CorteCaja, Gastos e Insumos."
    )
    
    add_code_table_with_explanation(
        doc,
        title="Código Implementado por Yadier Pech Tun — Precarga de Catálogo y Etiquetas Base (App/BaseDatos/DatabaseInitializer.cs - Método SeedInitialData)",
        code_text=extract_method_exact("BaseDatos/DatabaseInitializer.cs", "SeedInitialData"),
        hito_explanation="Este método resuelve el hito de 'SQLite con reglas base y etiquetas precargadas'. Verifica si la base de datos se encuentra vacía al instalar el sistema y, de ser así, inserta automáticamente la cuenta de usuario `admin` con su hash de seguridad sha256 y un catálogo de 8 servicios de lavado y secado listos para operar desde el primer minuto."
    )
    
    add_code_table_with_explanation(
        doc,
        title="Código Implementado por Yadier Pech Tun — Repositorio Relacional de Partidas (App/Repositorios/DetallePedidoRepositorio.cs)",
        code_text=read_file_exact("Repositorios/DetallePedidoRepositorio.cs"),
        hito_explanation="Este hito arquitectónico permite persistir y consultar las líneas o partidas individuales de cada pedido en SQL. Utiliza parámetros transaccionales (`@idPedido`, `@idServicio`, `@cantidad`, etc.) para prevenir inyecciones SQL y garantizar operaciones atómicas ultrarrápidas."
    )
    
    add_code_table_with_explanation(
        doc,
        title="Código Implementado por Jesus Leyva Chan — Repositorio SQL del Catálogo de Servicios (App/Repositorios/ServicioRepositorio.cs)",
        code_text=read_file_exact("Repositorios/ServicioRepositorio.cs"),
        hito_explanation="Este repositorio implementa el acceso completo a los datos del catálogo comercial. Permite extraer la lista completa de servicios activos por orden alfabético y persistir de forma segura nuevos precios o ajustes en `lavanderia.db`."
    )
    
    add_code_table_with_explanation(
        doc,
        title="Código Implementado por Jesus Leyva Chan — Repositorio SQL de Clientes e Historial (App/Repositorios/ClienteRepositorio.cs)",
        code_text=read_file_exact("Repositorios/ClienteRepositorio.cs"),
        hito_explanation="Este hito completa la arquitectura de repositorios para clientes. Encapsula las sentencias SQL de alta, búsqueda instantánea por nombre o teléfono (`LIKE @query%`) y actualización relacional de visitas acumuladas para el sistema de lealtad."
    )

    # =========================================================================
    # ACTIVIDAD 5
    # =========================================================================
    add_header_p(doc, "6. Actividad: Aprobación del prototipo y cierre de diseño", level=1)
    add_activity_table(
        doc,
        act_name="Aprobación del prototipo y cierre de diseño",
        f_init="2026-06-12",
        f_end="2026-06-19",
        duration="7 días",
        resp="Todos",
        dep="Diseño de interfaz y flujo en Figma",
        hitos="Validación visual, ajustes finales y aprobación para desarrollo",
        avance="100%",
        nota="Esta tarjeta marca el punto de cierre de paso entre diseño y construcción.\nAVANCE TOTAL 50%"
    )
    
    add_header_p(doc, "¿Qué se pretende hacer a rasgos generales en esta actividad?", level=2)
    add_body_p(doc, "En este punto de inflexión y cierre formal entre el diseño de experiencia y la programación intensiva, se pretende validar de forma integral todas las pantallas de la plataforma y aplicar los ajustes finales de formato e interacción. Específicamente, se busca consolidar en el frontend la pantalla completa de configuración visual y de negocio, estructurar en backend el modelo digital del comprobante (ticket) que se entregará al cliente final de acuerdo a la validación de diseño, e implementar un servicio centralizado de cuadros de diálogo y confirmación modal que sustituya las alertas básicas del navegador o sistema operativo por una estética corporativa.")
    
    add_header_p(doc, "Explicación de Cada Integrante con su Respectivo Rol", level=2)
    add_body_p(doc, "Como responsable de Frontend, Daniel Moo ajustó, pulió y finalizó la vista de administración de parámetros estéticos y operativos del sistema en Configuraciones.razor. Esta pantalla permite a la administración seleccionar el tema visual (Oscuro, Monocrómico, etc.) y editar los datos de cabecera que se muestran en la plataforma, reflejando fielmente la estructura aprobada por el usuario en Figma.", bold_prefix="Daniel Moo (Rol: Frontend): ")
    add_body_p(doc, "Como responsable del área de Backend & Base de Datos, Yadier Pech Tun construyó la entidad transaccional de comprobantes Ticket.cs, donde programó el layout digital exacto del recibo comercial: cabecera institucional, fecha, número de folio, desglose por kilogramo/pieza y totales financieros listos para enviarse a impresión o PDF según el diseño aprobado.", bold_prefix="Yadier Pech Tun (Rol: Backend & Base de Datos): ")
    add_body_p(doc, "Como responsable de Backend, Leyva Chan desarrolló el servicio de interrupción y diálogo institucional CustomMessageBox.cs. Este servicio estático permite que cualquier módulo o página del sistema muestre confirmaciones de seguridad y mensajes de error o éxito con el estilo visual del diseño aprobado, ejecutando callbacks asíncronos para acciones del mostrador.", bold_prefix="Jesus Leyva Chan (Rol: Backend): ")
    
    add_header_p(doc, "¿Qué y Cómo se Implementó en el Código? (Referencia Verbatim y Explicación por Hito)", level=2)
    
    add_code_table_with_explanation(
        doc,
        title="Código Implementado por Daniel Moo — Vista Integral de Configuración Visual y del Negocio (App/Pages/Configuraciones.razor)",
        code_text=read_file_exact("Pages/Configuraciones.razor"),
        hito_explanation="Esta interfaz Razor completa representa el hito de 'Ajustes finales y validación visual del módulo de configuración'. Permite al personal directivo modificar en tiempo real el tema de color del POS, los datos de contacto y el fondo inicial de caja, aplicando de forma reactiva el diseño validado al resto de las pantallas del sistema."
    )
    
    add_code_table_with_explanation(
        doc,
        title="Código Implementado por Yadier Pech Tun — Estructura de Comprobante del Prototipo Validado (App/Modelos/Ticket.cs)",
        code_text=read_file_exact("Modelos/Ticket.cs"),
        hito_explanation="Este hito materializa la validación visual de la salida del negocio: el ticket de mostrador. La clase Ticket organiza de forma estructurada los datos del pedido y el cliente, ofreciendo un método auxiliar para formatear la salida en texto monoespaciado lista para impresoras térmicas o exportación digital con cabeceras formales."
    )
    
    add_code_table_with_explanation(
        doc,
        title="Código Implementado por Jesus Leyva Chan — Servicio Institucional de Cuadros de Diálogo Modales (App/Servicios/CustomMessageBox.cs)",
        code_text=read_file_exact("Servicios/CustomMessageBox.cs"),
        hito_explanation="Este servicio cumple con el hito de 'Aprobación para desarrollo y estandarización de diálogos de usuario'. Al encapsular la invocación de ventanas emergentes en C#, garantiza que todo mensaje de confirmación, advertencia o error en el sistema mantenga una presentación uniforme y ergonómica, respetando el cierre de diseño."
    )

    # =========================================================================
    # ACTIVIDAD 6
    # =========================================================================
    add_header_p(doc, "7. Actividad: Backend: autenticación y usuarios", level=1)
    add_activity_table(
        doc,
        act_name="Backend: autenticación y usuarios",
        f_init="2026-06-23",
        f_end="2026-07-04",
        duration="11 días",
        resp="Yadier Pech Tun y Leyva Chan",
        dep="definición de módulos y backlog",
        hitos="Inicio de sesión(Leyva), roles (Yadier), sesión segura(Leyva) y gestión básica de usuarios(Yadier)",
        avance="100%",
        nota="Este módulo debe quedar estable antes de avanzar al resto de la lógica del sistema.\nAVANCE TOTAL AL TERMINAR: 60%"
    )
    
    add_header_p(doc, "¿Qué se pretende hacer a rasgos generales en esta actividad?", level=2)
    add_body_p(doc, "El propósito fundamental en este apartado es blindar y estabilizar por completo la seguridad en el backend antes de permitir el procesamiento de órdenes en mostrador. Se pretende implementar la validación de inicio de sesión contra SQL (`Login`), la apertura y mantenimiento de una sesión global en memoria (`sesión segura`), la tipificación de privilegios jerárquicos entre Operador y Administrador (`roles`) y el motor de operaciones relacionales para el alta, búsqueda y actualización de cuentas de personal (`gestión básica de usuarios`).")
    
    add_header_p(doc, "Explicación de Cada Integrante con su Respectivo Rol", level=2)
    add_body_p(doc, "En cumplimiento puntual y exacto con los hitos asignados ('Inicio de sesión y sesión segura') bajo su rol en Backend, Leyva Chan construyó el servicio transaccional LoginServicio.cs. Este motor consulta a la base de datos SQL el usuario proporcionado, verifica de forma criptográfica la coincidencia de contraseña en SHA256 o texto seguro, revisa que la cuenta no haya sido desactivada (`Activo == true`) y, upon success, invoca a SessionManager.SetUser() para asegurar la sesión del usuario en memoria antes de permitir el ingreso al POS.", bold_prefix="Jesus Leyva Chan (Rol: Backend - Hitos: Inicio de sesión y sesión segura): ")
    add_body_p(doc, "En cumplimiento puntual con los hitos asignados ('Roles y gestión básica de usuarios') en su rol de Backend & Base de Datos, Yadier Pech Tun desarrolló tres componentes relacionales separados e independientes:\n"
                    "1) La entidad relacional Usuario.cs, que mapea la tabla de base de datos e incluye la asignación transaccional del rol y estado de la cuenta.\n"
                    "2) El modelo derivado Admin.cs, especializado para facultades directivas.\n"
                    "3) El repositorio SQL completo UsuarioRepositorio.cs, encargado de ejecutar las sentencias DML (INSERT, UPDATE, SELECT por ID y verificación de duplicados de nombre de usuario) contra la base de datos.", bold_prefix="Yadier Pech Tun (Rol: Backend & Base de Datos - Hitos: Roles y gestión básica de usuarios): ")
    add_body_p(doc, "Dado que esta tarjeta es por definición una fase intensiva de backend, autenticación criptográfica, sesiones en memoria y repositorios SQL para credenciales, Daniel Moo (Frontend) interviene en el consumo posterior de estos servicios desde las pantallas de login y administración de usuarios.", bold_prefix="Daniel Moo (Rol: Frontend): ")
    
    add_header_p(doc, "¿Qué y Cómo se Implementó en el Código? (Referencia Verbatim y Explicación por Hito)", level=2)
    
    add_code_table_with_explanation(
        doc,
        title="Código Implementado por Jesus Leyva Chan — Servicio Transaccional de Autenticación (App/Servicios/LoginServicio.cs)",
        code_text=read_file_exact("Servicios/LoginServicio.cs"),
        hito_explanation="Este servicio cumple con los hitos de 'Inicio de sesión' y 'Sesión segura'. Ejecuta una consulta relacional sobre la tabla Usuarios, verifica que la contraseña coincida exactamente con el hash registrado, valida que la cuenta se encuentre activa e invoca inmediatamente al gestor de sesión para consolidar el ingreso y desbloquear la interfaz."
    )
    
    add_code_table_with_explanation(
        doc,
        title="Código Implementado por Yadier Pech Tun — Entidad Relacional de Usuario y Tipificación de Roles (App/Modelos/Usuario.cs)",
        code_text=read_file_exact("Modelos/Usuario.cs"),
        hito_explanation="Este archivo satisface la primera parte del hito 'Roles (Yadier)'. La clase Usuario modela los atributos de acceso de cada empleado en `lavanderia.db`, estableciendo de manera explícita si la cuenta cuenta con el privilegio transaccional `Rol = 'Admin'` u `'Operador'`."
    )
    
    add_code_table_with_explanation(
        doc,
        title="Código Implementado por Yadier Pech Tun — Especialización Criptográfica y de Privilegios para Administradores (App/Modelos/Admin.cs)",
        code_text=read_file_exact("Modelos/Admin.cs"),
        hito_explanation="Este hito profundiza en el control de roles del sistema al crear una especialización de dominio para directivos. Admin hereda los comportamientos de Usuario pero fuerza desde su construcción que el rol sea invariablemente 'Admin', proveyendo métodos auxiliares de validación criptográfica."
    )
    
    add_code_table_with_explanation(
        doc,
        title="Código Implementado por Yadier Pech Tun — Repositorio SQL para la Gestión Básica de Usuarios (App/Repositorios/UsuarioRepositorio.cs)",
        code_text=read_file_exact("Repositorios/UsuarioRepositorio.cs"),
        hito_explanation="Este repositorio cumple al 100% con el hito 'Gestión básica de usuarios (Yadier)'. Encapsula las sentencias SQL parametrizadas para registrar nuevos operadores, actualizar contraseñas o roles, extraer el listado completo de personal y verificar la existencia previa de un nombre de usuario para evitar duplicidades en SQL."
    )

    # =========================================================================
    # ACTIVIDAD 7
    # =========================================================================
    add_header_p(doc, "8. Actividad: Backend: clientes, servicios y órdenes", level=1)
    add_activity_table(
        doc,
        act_name="Backend: clientes, servicios y órdenes",
        f_init="2026-07-05",
        f_end="2026-07-14",
        duration="5 días",
        resp="Yadier Pech Tun y Leyva Chan",
        dep="Autenticación y usuarios listos",
        hitos="CRUD principal(Yadier), registro de clientes (Leyva)",
        avance="100%",
        nota="Aquí se concentra la lógica más importante de la lavandería.\nAVANCE TOTAL AL TERMINAR: 65%"
    )
    
    add_header_p(doc, "¿Qué se pretende hacer a rasgos generales en esta actividad?", level=2)
    add_body_p(doc, "En este apartado, considerado el corazón operativo de la lavandería, se pretende orquestar toda la lógica transaccional del negocio. El objetivo es que el sistema sea capaz de registrar clientes y calcular su lealtad por acumulación de visitas (`registro de clientes`), así como ejecutar el `CRUD principal`: procesar de forma atómica el alta, modificación y seguimiento de órdenes de pedido, y —de manera fundamental— automatizar la deducción de inventario de insumos (detergentes, suavizantes, cloro) restando mililitros o gramos del almacén en función exacta de la fórmula por cada kilogramo de ropa procesada en la orden.")
    
    add_header_p(doc, "Explicación de Cada Integrante con su Respectivo Rol", level=2)
    add_body_p(doc, "En cumplimiento con el hito de 'CRUD principal' asignado y como titular de Backend y Base de Datos, Yadier Pech Tun desarrolló el motor transaccional más completo y robusto de la aplicación en tres piezas separadas:\n"
                    "1) El servicio orquestador PedidoServicio.cs, que procesa la creación integral del pedido, asigna folios, calcula montos financieros e invoca la actualización en base de datos.\n"
                    "2) El repositorio relacional PedidoRepositorio.cs, que ejecuta transacciones SQL multitabla (encabezado y detalles) con manejo de rollbacks en caso de fallo.\n"
                    "3) El motor de deducción automática InventarioAutomatizacion.cs, que consulta las recetas de insumos del catálogo por cada partida de la orden y descuenta de manera automática los mililitros o gramos exactos del inventario SQL al procesarse una carga.", bold_prefix="Yadier Pech Tun (Rol: Backend & Base de Datos - Hito: CRUD principal): ")
    add_body_p(doc, "En cumplimiento con el hito de 'registro de clientes' asignado en su rol de Backend, Leyva Chan programó la lógica de negocio y persistencia de clientes en dos componentes separados:\n"
                    "1) ClienteServicio.cs, que centraliza la validación, alta en mostrador, actualización y búsqueda rápida por coincidencia fonética o telefónica.\n"
                    "2) Cliente.cs, la entidad relacional que almacena el contacto, acumula visitas e indica de forma dinámica cuándo un cliente regular se hace acreedor a promociones o beneficios de lealtad.", bold_prefix="Jesus Leyva Chan (Rol: Backend - Hito: registro de clientes): ")
    add_body_p(doc, "Al tratarse de una fase puramente centrada en servicios transaccionales de backend, lógica de inventario automatizado por fórmulas y repositorios relacionales para pedidos y clientes, Daniel Moo (Frontend) se encarga posteriormente de conectar estos potentes motores a los componentes visuales de cobro y registro en mostrador.", bold_prefix="Daniel Moo (Rol: Frontend): ")
    
    add_header_p(doc, "¿Qué y Cómo se Implementó en el Código? (Referencia Verbatim y Explicación por Hito)", level=2)
    
    add_code_table_with_explanation(
        doc,
        title="Código Implementado por Yadier Pech Tun — Servicio Orquestador del CRUD Principal de Pedidos (App/Servicios/PedidoServicio.cs)",
        code_text=read_file_exact("Servicios/PedidoServicio.cs"),
        hito_explanation="Este servicio central cumple con el 'CRUD principal (Yadier)'. Coordina la creación y modificación de pedidos en la capa de negocio, validando que existan partidas activas en la orden, calculando importes totales e invocando al repositorio transaccional y al motor de deducción de inventario."
    )
    
    add_code_table_with_explanation(
        doc,
        title="Código Implementado por Yadier Pech Tun — Repositorio SQL Transaccional Multitabla de Pedidos (App/Repositorios/PedidoRepositorio.cs)",
        code_text=read_file_exact("Repositorios/PedidoRepositorio.cs"),
        hito_explanation="Este repositorio materializa la persistencia relacional del CRUD principal. Utiliza transacciones explícitas de SQLite (`BeginTransaction()`) para insertar o actualizar en un solo paso atómico tanto la cabecera del pedido (`Pedidos`) como cada uno de sus renglones (`DetallesPedido`), garantizando cero pérdidas de información."
    )
    
    add_code_table_with_explanation(
        doc,
        title="Código Implementado por Yadier Pech Tun — Motor Transaccional de Deducción Automática de Inventario (App/Servicios/InventarioAutomatizacion.cs)",
        code_text=read_file_exact("Servicios/InventarioAutomatizacion.cs"),
        hito_explanation="Este hito técnico de alto nivel corona el CRUD principal al conectar los pedidos con el inventario. La clase InventarioAutomatizacion recorre cada partida de la orden entrante, extrae la receta técnica del servicio (cuántos ml/gr de detergente o suavizante se consumen por cada kg de ropa) y actualiza automáticamente las existencias en `lavanderia.db`, evitando desabastos en la lavandería."
    )
    
    add_code_table_with_explanation(
        doc,
        title="Código Implementado por Jesus Leyva Chan — Servicio Transaccional del Registro y Lealtad de Clientes (App/Servicios/ClienteServicio.cs)",
        code_text=read_file_exact("Servicios/ClienteServicio.cs"),
        hito_explanation="Este servicio cumple al 100% con el hito 'registro de clientes (Leyva)'. Expone métodos limpios para dar de alta o editar clientes en mostrador, buscar por teléfono o nombre al instante e incrementar en una visita el contador de lealtad cada vez que el cliente liquida una orden en el mostrador."
    )
    
    add_code_table_with_explanation(
        doc,
        title="Código Implementado por Jesus Leyva Chan — Entidad Relacional y Propiedades de Fidelidad de Clientes (App/Modelos/Cliente.cs)",
        code_text=read_file_exact("Modelos/Cliente.cs"),
        hito_explanation="Esta entidad completa el registro transaccional de clientes. Almacena la identidad relacional, teléfono, dirección y un contador de `VisitasAcumuladas`, ofreciendo una propiedad dinámica que determina cuándo un cliente frecuente alcanza el umbral para beneficios de fidelidad."
    )

    # =========================================================================
    # ACTIVIDAD 8
    # =========================================================================
    add_header_p(doc, "9. Actividad: Frontend: pantallas principales", level=1)
    add_activity_table(
        doc,
        act_name="Frontend: pantallas principales",
        f_init="2026-07-10",
        f_end="2026-07-17",
        duration="7 días",
        resp="Daniel Moo",
        dep="Prototipo de Figma aprobado",
        hitos="Login, inicio, clientes, órdenes, navegación y formularios principales",
        avance="100%",
        nota="Debe respetar el diseño aprobado para evitar retrabajo en la integración.\nAVANCE TOTAL AL TERMINAR: 80%"
    )
    
    add_header_p(doc, "¿Qué se pretende hacer a rasgos generales en esta actividad?", level=2)
    add_body_p(doc, "En esta actividad, bajo responsabilidad exclusiva del área de Frontend, se pretende construir y diagramar todas las pantallas transaccionales e interactivas de la plataforma, respetando con máxima fidelidad la estética, colores y flujos validados en Figma. El objetivo es proporcionar al personal de mostrador interfaces visuales limpias, rápidas y responsivas en Blazor Hybrid / Razor para el ingreso al sistema (`Login`), cobro ágil en punto de venta (`Cobro`), monitoreo interactivo de lavadoras en tiempo real (`Maquinas`), directorio y fidelidad de clientes (`Clientes`) y centro de control operativo de pedidos (`Pedidos`).")
    
    add_header_p(doc, "Explicación de Cada Integrante con su Respectivo Rol", level=2)
    add_body_p(doc, "Al ser el responsable único y exclusivo de la actividad en su rol de Frontend, Daniel Moo desarrolló la totalidad de las pantallas principales en componentes Razor funcionales. Para respetar el mandato de no agrupar archivos o pantallas distintas en un solo bloque, se presentan y explican de manera independiente cada una de las 5 vistas transaccionales maestras construidas por Daniel Moo:\n"
                    "1) Login.razor: Pantalla de autenticación con captura reactiva de credenciales y manejo visual de errores.\n"
                    "2) Cobro.razor: Terminal de Punto de Venta (POS) donde el operador visualiza las prendas del carrito, elige entre cobro con efectivo o tarjeta y calcula en tiempo real el cambio a entregar al cliente.\n"
                    "3) Maquinas.razor: Tablero visual dinámico que representa en cuadrícula el estado del parque de lavadoras y secadoras, cambiando de color (disponible vs ocupada) e informando el tiempo restante por ciclo.\n"
                    "4) Clientes.razor: Interfaz de mostrador con buscador instantáneo, alta emergente de nuevos contactos y visualización de visitas y lealtad.\n"
                    "5) Pedidos.razor: Centro de seguimiento operativo con filtrado por estados y actualización visual interactiva.", bold_prefix="Daniel Moo (Rol: Frontend - Responsable Único): ")
    add_body_p(doc, "En esta actividad de diseño e implementación visual de vistas en Blazor/Razor, Yadier Pech Tun (Base de Datos/Backend) y Leyva Chan (Backend) no intervienen de forma directa, dado que sus motores de servicios y repositorios fueron finalizados en las etapas previas para que el frontend de Daniel Moo los consuma con total solidez.", bold_prefix="Yadier Pech Tun y Jesus Leyva Chan: ")
    
    add_header_p(doc, "¿Qué y Cómo se Implementó en el Código? (Referencia Verbatim y Explicación por Hito)", level=2)
    
    add_code_table_with_explanation(
        doc,
        title="Código Implementado por Daniel Moo — Pantalla de Ingreso y Autenticación POS (App/Pages/Login.razor)",
        code_text=read_file_exact("Pages/Login.razor"),
        hito_explanation="Este componente Razor cumple con el hito 'Pantalla de Login y acceso al sistema'. Presenta un formulario visual ergonómico con validación en tiempo real que invoca al servicio de inicio de sesión, retroalimentando al operador ante credenciales incorrectas y abriendo la navegación en caso de éxito."
    )
    
    add_code_table_with_explanation(
        doc,
        title="Código Implementado por Daniel Moo — Terminal de Cobro POS y Cálculo Visual de Cambio (App/Pages/Cobro.razor)",
        code_text=read_file_exact("Pages/Cobro.razor"),
        hito_explanation="Este componente es uno de los hitos visuales más importantes: el terminal de cobro en mostrador. Muestra el desglose completo del pedido, permite seleccionar el método de pago con botones interactivos de alta visibilidad y calcula matemáticamente al instante el cambio exacto a devolver según el billete recibido."
    )
    
    add_code_table_with_explanation(
        doc,
        title="Código Implementado por Daniel Moo — Tablero Visual del Estado y Temporizadores de Máquinas (App/Pages/Maquinas.razor)",
        code_text=read_file_exact("Pages/Maquinas.razor"),
        hito_explanation="Este componente resuelve el hito de 'Navegación e interfaz de control de lavadoras'. Representa en tarjetas interactivas cada equipo de la lavandería, asignando colores intuitivos (verde para disponible, rojo/naranja en ciclo) y gestionando temporizadores en vivo para optimizar el flujo de trabajo."
    )
    
    add_code_table_with_explanation(
        doc,
        title="Código Implementado por Daniel Moo — Directorio Visual e Interactivo de Clientes (App/Pages/Clientes.razor)",
        code_text=read_file_exact("Pages/Clientes.razor"),
        hito_explanation="Esta interfaz satisface el hito de 'Pantalla principal de Clientes y formularios'. Permite al personal de mostrador buscar contactos por teclado al instante, registrar nuevos clientes mediante un modal emergente limpio y monitorear los puntos de lealtad acumulados por cada usuario."
    )
    
    add_code_table_with_explanation(
        doc,
        title="Código Implementado por Daniel Moo — Centro Maestro Visual de Gestión de Pedidos (App/Pages/Pedidos.razor)",
        code_text=read_file_exact("Pages/Pedidos.razor"),
        hito_explanation="Esta vista completa el hito de 'Pantallas principales y seguimiento de órdenes'. Organiza todos los pedidos activos o históricos en tablas y tarjetas con filtrado instantáneo por estado operativo, permitiendo al mostrador avanzar las prendas al siguiente paso con un solo clic."
    )

    # =========================================================================
    # ACTIVIDAD 9
    # =========================================================================
    add_header_p(doc, "10. Actividad: Integración frontend + backend", level=1)
    add_activity_table(
        doc,
        act_name="Integración frontend + backend",
        f_init="2026-07-10",
        f_end="2026-07-17",
        duration="7 días",
        resp="Todos",
        dep="Backend y frontend listos y aprobación del usuario",
        hitos="Conexión total, validación de flujos y pruebas de integración",
        avance="100%",
        nota="Si aparece un error aquí, conviene corregirlo antes de entrar a pruebas finales.\nAVANCE TOTAL 85%"
    )
    
    add_header_p(doc, "¿Qué se pretende hacer a rasgos generales en esta actividad?", level=2)
    add_body_p(doc, "En esta fase decisiva de integración se pretende interconectar y sincronizar de forma absoluta todas las vistas gráficas del frontend en Blazor con los servicios, motores e infraestructura relacional de C# y SQLite. El objetivo es que la aplicación arranque inicializando de manera limpia la base de datos e inyectando las dependencias transaccionales (`App.xaml.cs`), y que cada acción o clic en el mostrador (como cobrar una orden o cambiar de estado un pedido) ejecute transacciones reales end-to-end con validación de flujos.")
    
    add_header_p(doc, "Explicación de Cada Integrante con su Respectivo Rol", level=2)
    add_body_p(doc, "Desde su responsabilidad de Backend y Base de Datos, Yadier Pech Tun construyó el punto central de arranque e inyección de dependencias en App.xaml.cs. En este módulo, Yadier vinculó el ciclo de vida de la interfaz gráfica con el motor relacional, ejecutando de manera automática la verificación, creación (`Initialize()`) y precarga (`SeedInitialData()`) de la base de datos SQLite antes de renderizar la primera ventana del POS.", bold_prefix="Yadier Pech Tun (Rol: Backend & Base de Datos): ")
    add_body_p(doc, "Desde el área de Backend, Leyva Chan desarrolló la integración operativa y transaccional para el procesamiento real de cobros en la capa de negocio: PagoServicio.cs. Este servicio conecta de manera directa la terminal visual de cobro del mostrador con el repositorio SQL de pedidos, asentando el pago, actualizando el saldo pendiente del cliente y modificando el estado del pedido a pagado.", bold_prefix="Jesus Leyva Chan (Rol: Backend): ")
    add_body_p(doc, "En su rol de Frontend, Daniel Moo interconectó las pantallas emergentes con los servicios backend inyectados. En específico, programó el componente de transición operativa CambiarEstado.razor, que permite modificar desde el mostrador el estado de una orden en base de datos, y estructuró la importación global de namespaces y servicios en _Imports.razor para que todos los componentes tengan acceso sin restricciones a los motores relacionales.", bold_prefix="Daniel Moo (Rol: Frontend): ")
    
    add_header_p(doc, "¿Qué y Cómo se Implementó en el Código? (Referencia Verbatim y Explicación por Hito)", level=2)
    
    add_code_table_with_explanation(
        doc,
        title="Código Implementado por Yadier Pech Tun — Contenedor de Arranque e Inicialización de Base de Datos (App/App.xaml.cs)",
        code_text=read_file_exact("App.xaml.cs"),
        hito_explanation="Este archivo cumple con el hito 'Conexión total y arranque transaccional del sistema'. Al iniciarse la aplicación WPF/Blazor, `OnStartup` invoca inmediatamente al inicializador y precargador de base de datos de Yadier (`DatabaseInitializer`), asegurando que el motor SQLite esté activo y validado antes de que el usuario vea la pantalla."
    )
    
    add_code_table_with_explanation(
        doc,
        title="Código Implementado por Jesus Leyva Chan — Servicio Transaccional de Integración de Pagos y Cobros (App/Servicios/PagoServicio.cs)",
        code_text=read_file_exact("Servicios/PagoServicio.cs"),
        hito_explanation="Este servicio satisface el hito de 'Validación de flujos de cobro en la capa de integración'. Recibe los parámetros de pago capturados en la interfaz visual, ejecuta la inserción relacional en SQL y actualiza de manera atómica el saldo adeudado y el estado financiero del pedido en la base de datos."
    )
    
    add_code_table_with_explanation(
        doc,
        title="Código Implementado por Daniel Moo — Componente Integrado para el Cambio Transaccional de Estados (App/Pages/CambiarEstado.razor)",
        code_text=read_file_exact("Pages/CambiarEstado.razor"),
        hito_explanation="Este componente cumple con el hito de 'Pruebas de integración visuales y operativas'. Conecta el clic del operador en mostrador directamente con la actualización relacional en SQL del pedido, moviendo las prendas desde `Lavando` hasta `Listo` o `Entregado` con retroalimentación en vivo."
    )
    
    add_code_table_with_explanation(
        doc,
        title="Código Implementado por Daniel Moo — Registro Global de Dependencias y Namespaces del Frontend (App/_Imports.razor)",
        code_text=read_file_exact("_Imports.razor"),
        hito_explanation="Este archivo representa el puente estructural del frontend hacia la arquitectura inyectada. Al declarar de manera global los using hacia `Modelos`, `Servicios` y `Repositorios`, permite que cualquier componente visual de Blazor consuma las entidades de base de datos de manera directa y segura."
    )

    # =========================================================================
    # ACTIVIDAD 10
    # =========================================================================
    add_header_p(doc, "11. Actividad: Pruebas funcionales y corrección de errores", level=1)
    add_activity_table(
        doc,
        act_name="Pruebas funcionales y corrección de errores",
        f_init="2026-07-15",
        f_end="2026-07-19",
        duration="4 días",
        resp="Todos",
        dep="Integración completa de frontend y backend",
        hitos="Casos de prueba, revisión de fallos, ajustes de interfaz y validación de procesos",
        avance="100%",
        nota="Conviene registrar los errores por prioridad para resolver primero los que bloquean la entrega.\nAVANCE TOTAL 95%"
    )
    
    add_header_p(doc, "¿Qué se pretende hacer a rasgos generales en esta actividad?", level=2)
    add_body_p(doc, "En esta etapa de aseguramiento de calidad y estabilización, se pretende ejecutar pruebas funcionales exhaustivas y resolver todo error o anomalía antes del despliegue en el entorno de producción. Se busca instrumentar una bitácora transaccional de auditoría en SQL que registre e historice cada movimiento y posible fallo operativo (`casos de prueba y revisión de fallos`), habilitar un sistema relacional de notificaciones y alertas en campana (`validación de procesos`), y comprobar que el tablero ejecutivo principal (`Dashboard`) agregue métricas en vivo con tolerancia a fallos.")
    
    add_header_p(doc, "Explicación de Cada Integrante con su Respectivo Rol", level=2)
    add_body_p(doc, "Como titular de Backend y Base de Datos, Yadier Pech Tun implementó la infraestructura de observabilidad y trazabilidad operativa del sistema en dos clases separadas:\n"
                    "1) La entidad Auditoria.cs, que modela cada evento del sistema registrando fecha exacta, usuario responsable, acción, tabla modificada y descripción técnica del suceso.\n"
                    "2) El repositorio AuditoriaRepositorio.cs, que inserta y consulta de forma inmutable estos registros en `lavanderia.db`, permitiendo clasificar incidencias y auditar el comportamiento del personal en las pruebas.", bold_prefix="Yadier Pech Tun (Rol: Backend & Base de Datos): ")
    add_body_p(doc, "Desde su rol de Backend, Leyva Chan construyó el sistema de notificaciones y alertas transaccionales del mostrador en dos componentes independientes:\n"
                    "1) La entidad Notificacion.cs, que estructura los avisos automáticos de órdenes listas o incidencias en lavadoras.\n"
                    "2) El repositorio NotificacionRepositorio.cs, que persiste los avisos en SQL y gestiona transaccionalmente el marcado de alertas leídas o pendientes en la campana de notificaciones.", bold_prefix="Jesus Leyva Chan (Rol: Backend): ")
    add_body_p(doc, "Desde el área Frontend, Daniel Moo ejecutó las pruebas visuales y construyó el Tablero Ejecutivo en Vivo Dashboard.razor. Este componente recopila las métricas agregadas desde los repositorios (ingresos del día, órdenes en espera, máquinas en uso) y las presenta en tarjetas gráficas con bloques `try-catch` que toleran y recuperan cualquier error transaccional sin interrumpir la visualización.", bold_prefix="Daniel Moo (Rol: Frontend): ")
    
    add_header_p(doc, "¿Qué y Cómo se Implementó en el Código? (Referencia Verbatim y Explicación por Hito)", level=2)
    
    add_code_table_with_explanation(
        doc,
        title="Código Implementado por Yadier Pech Tun — Entidad Transaccional de Auditoría y Trazabilidad (App/Modelos/Auditoria.cs)",
        code_text=read_file_exact("Modelos/Auditoria.cs"),
        hito_explanation="Esta entidad cumple con el hito 'Revisión de fallos y bitácora de auditoría transaccional'. Permite que el sistema registre de manera inmutable quién hizo qué y cuándo en la base de datos, facilitando la identificación e investigación de errores durante las pruebas funcionales."
    )
    
    add_code_table_with_explanation(
        doc,
        title="Código Implementado por Yadier Pech Tun — Repositorio SQL para el Registro e Historial de Auditoría (App/Repositorios/AuditoriaRepositorio.cs)",
        code_text=read_file_exact("Repositorios/AuditoriaRepositorio.cs"),
        hito_explanation="Este repositorio ejecuta la persistencia de los casos de auditoría en `lavanderia.db`. Provee sentencias parametrizadas ultrarrápidas para insertar eventos operativos en tiempo real y recuperar historiales cronológicos de validación de procesos."
    )
    
    add_code_table_with_explanation(
        doc,
        title="Código Implementado por Jesus Leyva Chan — Entidad Transaccional de Notificaciones Operativas (App/Modelos/Notificacion.cs)",
        code_text=read_file_exact("Modelos/Notificacion.cs"),
        hito_explanation="Este hito materializa la 'Validación de procesos y alertas en campana'. La clase Notificacion define los avisos operativos que informan al personal sobre órdenes concluidas o anomalías, incluyendo una propiedad transaccional para controlar si el aviso ya fue leído."
    )
    
    add_code_table_with_explanation(
        doc,
        title="Código Implementado por Jesus Leyva Chan — Repositorio SQL de Alertas y Notificaciones (App/Repositorios/NotificacionRepositorio.cs)",
        code_text=read_file_exact("Repositorios/NotificacionRepositorio.cs"),
        hito_explanation="Este repositorio administra la campana de notificaciones en SQL. Consulta los avisos pendientes del usuario actual y actualiza el estado transaccional de lectura una vez que el operador revisa la alerta en la pantalla."
    )
    
    add_code_table_with_explanation(
        doc,
        title="Código Implementado por Daniel Moo — Tablero Ejecutivo en Vivo con Métricas y Tolerancia a Fallos (App/Pages/Dashboard.razor)",
        code_text=read_file_exact("Pages/Dashboard.razor"),
        hito_explanation="Este componente cumple con el hito 'Ajustes de interfaz, métricas operativas y validación de estabilización'. Renderiza el centro de mando del negocio con resúmenes estadísticos en tiempo real, encapsulando las consultas en bloques de seguridad que garantizan cero caídas en mostrador."
    )

    # =========================================================================
    # ACTIVIDAD 11
    # =========================================================================
    add_header_p(doc, "12. Actividad: Documentación final y entrega", level=1)
    add_activity_table(
        doc,
        act_name="Documentación final y entrega",
        f_init="2026-07-18",
        f_end="2026-07-22",
        duration="4 días",
        resp="Todos",
        dep="Pruebas aprobadas",
        hitos="Documentación técnica, presentación final, cierre del proyecto y enlace de Figma adjunto",
        avance="0%",
        nota="Aquí se deja lista la versión final del proyecto para revisión y entrega.\nAVANCE TOTAL 100%"
    )
    
    add_header_p(doc, "¿Qué se pretende hacer a rasgos generales en esta actividad?", level=2)
    add_body_p(doc, "En esta actividad final de cierre y entrega, se pretende consolidar la versión definitiva de la solución LavanderiaApp v0.1 para revisión directiva y entrega formal. A nivel técnico y contable, esto exige implementar las estructuras matemáticas rigurosas para los cortes de caja de turno y el registro de egresos de caja chica (`Cierre de proyecto y contabilidad`), construir en el frontend la interfaz visual de exportación y consulta de reportes operativos (`Reportes`), y emparejar el sistema con un motor oficial de generación de documentos en PDF (`generar_pdf_actividad.py`) capaz de emitir en formato imprimible los cortes contables e informes con cabeceras institucionales.")
    
    add_header_p(doc, "Explicación de Cada Integrante con su Respectivo Rol", level=2)
    add_body_p(doc, "Como responsable de Backend y Base de Datos, Yadier Pech Tun construyó los modelos contables y de auditoría financiera que permiten realizar el balance de cierre sin discrepancias. Para respetar el principio de separar componentes independientes, se exponen por separado:\n"
                    "1) CorteCaja.cs: Entidad contable que registra los montos reportados por el empleado vs los esperados por el sistema, calculando la diferencia y asentando el cierre de turno.\n"
                    "2) Gasto.cs: Entidad para asentar salidas justificadas de efectivo de la caja chica para compra de insumos de emergencia o servicios.", bold_prefix="Yadier Pech Tun (Rol: Backend & Base de Datos): ")
    add_body_p(doc, "Como responsable de Backend, Leyva Chan desarrolló el repositorio transaccional relacional CorteCajaRepositorio.cs. Su implementación almacena de forma inmutable cada arqueo diario en la base de datos `lavanderia.db`, asegurando que los cierres contables queden debidamente registrados para la revisión técnica y la presentación final.", bold_prefix="Jesus Leyva Chan (Rol: Backend): ")
    add_body_p(doc, "En su calidad de titular de Frontend, Daniel Moo desarrolló la interfaz visual para la emisión, filtrado por fechas y consulta de balances operativos en Reportes.razor, y además conectó e implementó el script oficial de generación de informes y reportes en PDF generar_pdf_actividad.py con ReportLab. Esta aportación dota al establecimiento de reportes profesionales listos para imprimirse y descargarse en la entrega final.", bold_prefix="Daniel Moo (Rol: Frontend): ")
    
    add_header_p(doc, "¿Qué y Cómo se Implementó en el Código? (Referencia Verbatim y Explicación por Hito)", level=2)
    
    add_code_table_with_explanation(
        doc,
        title="Código Implementado por Yadier Pech Tun — Entidad Contable de Arqueo y Corte de Turno (App/Modelos/CorteCaja.cs)",
        code_text=read_file_exact("Modelos/CorteCaja.cs"),
        hito_explanation="Esta entidad cumple con la contabilidad de cierre del proyecto. Almacena de forma rigurosa los montos en efectivo contados por el cajero frente a lo que el motor SQL calculó (`EfectivoEsperado`), dejando constancia inmutable de cualquier sobrante o faltante al final de la jornada."
    )
    
    add_code_table_with_explanation(
        doc,
        title="Código Implementado por Yadier Pech Tun — Entidad Contable de Egresos de Caja Chica (App/Modelos/Gasto.cs)",
        code_text=read_file_exact("Modelos/Gasto.cs"),
        hito_explanation="Este hito complementa el cierre contable de la lavandería al modelar las salidas de efectivo. La clase Gasto permite asentar cualquier erogación operativa justificada, restando ese monto de manera automática del efectivo esperado en el corte de caja."
    )
    
    add_code_table_with_explanation(
        doc,
        title="Código Implementado por Jesus Leyva Chan — Repositorio SQL para Cierres de Caja (App/Repositorios/CorteCajaRepositorio.cs)",
        code_text=read_file_exact("Repositorios/CorteCajaRepositorio.cs"),
        hito_explanation="Este repositorio asegura el hito de 'Cierre contable transaccional relacional'. Persiste y consulta los cortes en `lavanderia.db`, permitiendo a la administración consultar el historial de balances diarios y verificar la honestidad y exactitud de cada turno laboral."
    )
    
    add_code_table_with_explanation(
        doc,
        title="Código Implementado por Daniel Moo — Interfaz de Usuario para Emisión y Consulta de Reportes (App/Pages/Reportes.razor)",
        code_text=read_file_exact("Pages/Reportes.razor"),
        hito_explanation="Esta pantalla Razor cumple con el hito de 'Consulta e interfaz de reportes del proyecto'. Proporciona al administrador una vista limpia para filtrar ingresos y cortes por rangos de fecha, presentando tablas totalizadas y botones de descarga oficial."
    )
    
    add_code_table_with_explanation(
        doc,
        title="Código Implementado por Daniel Moo — Motor Oficial de Generación de Informes y Reportes Operativos en PDF (App/generar_pdf_actividad.py)",
        code_text=read_file_exact("generar_pdf_actividad.py"),
        hito_explanation="Este script corona la 'Documentación técnica y presentación final'. Utiliza la librería profesional ReportLab para generar documentos PDF paginados con cabeceras corporativas, tablas con fondos temáticos e índices de auditoría para entregar un producto final de calidad ejecutiva."
    )

    # Rutas de salida en App/Shared como solicitó explícitamente el usuario
    out_paths = [
        r"C:\Users\Yadie\RiderProjects\LavanderiaApp0.1\App\Shared\Informe_Desarrollo_LavanderiaApp_Final_Completo.docx",
        r"C:\Users\Yadie\RiderProjects\LavanderiaApp0.1\App\Shared\Informe_Desarrollo_LavanderiaApp_Adaptado_Final.docx",
        r"C:\Users\Yadie\Downloads\Informe_Desarrollo_LavanderiaApp_Final_Completo.docx"
    ]
    saved = []
    for out_path in out_paths:
        try:
            doc.save(out_path)
            saved.append(out_path)
            print(f"¡Guardado exitoso en: {out_path}!")
        except PermissionError:
            print(f"[Aviso] El archivo {out_path} está abierto en Microsoft Word. Se guardó en una ruta alternativa.")
        except Exception as e:
            print(f"Error al guardar en {out_path}: {e}")
            
    if not saved:
        print("No se pudo guardar el archivo en ninguna ubicación.")
    else:
        print("¡Generación exitosa del informe completo!")

if __name__ == "__main__":
    generate_full_report()
